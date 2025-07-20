import argparse
import os
import re
import sys
from enum import Enum
from pathlib import Path
from typing import Dict

import docx.oxml.shared
import markdown
from bs4 import BeautifulSoup
from bs4.element import PageElement as BS4_Element
from docx import Document as DOCX_Document
from docx.enum.style import WD_STYLE_TYPE as DOCX_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL as DOCX_CELL_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH as DOCX_PARAGRAPH_ALIGN
from docx.enum.text import WD_BREAK as DOCX_BREAK_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE as DOCX_REL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table as DOCX_Table
from docx.table import _Cell as DOCX_Cell
from docx.text.font import Font as DOCX_FONT
from docx.text.paragraph import Paragraph as DOCX_Paragraph
from docx.text.parfmt import ParagraphFormat as DOCX_ParagraphFormat
from docx.text.run import Run as DOCX_Run

SCRIPT_DIR = Path(__file__).parent

##############################
# Define some defaults at module level for better performance
##############################
DEFAULT_OUTPUT_DIR = SCRIPT_DIR.parent / "data"
DEFAULT_CONFIG_FILE = SCRIPT_DIR / "resume_config.yaml"
DOCX_EXTENSION = "docx"
PDF_EXTENSION = "pdf"
DEFAULT_OUTPUT_FORMAT = DOCX_EXTENSION
MD_LINK_PATTERN = re.compile(r"\[(.*?)\]\((.*?)\)")
URL_PATTERN = re.compile(r"https?://[^\s]+|www\.[^\s]+")
EMAIL_PATTERN = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")


##############################
# Configs
##############################
class JobSubsection(Enum):
    """Maps markdown subsection headings to their corresponding document headings

    Properties:
        markdown_heading_level (str): The HTML tag name (e.g., 'h5', 'h6') in markdown
        markdown_text_lower (str): The lowercase text content of the heading to match
        docx_heading (str): The text to use as a heading in the Word document
        separator (str): Optional separator to append after the heading (default: "")
        bold (bool): Whether the heading should be bold (default: True)
        italic (bool): Whether the heading should be italic (default: False)
    """

    KEY_SKILLS = ("h5", "key skills", "Key Skills", ":", True, False)
    SUMMARY = ("h5", "summary", "Summary", ":", True, False)
    INTERNAL = ("h5", "internal", "Internal", ":", True, False)
    PROJECT_CLIENT = ("h5", "project/client", "Project/Client", ": ", True, True)
    RESPONSIBILITIES = (
        "h6",
        "responsibilities overview",
        "Responsibilities",
        ":",
        True,
        False,
    )
    ADDITIONAL_DETAILS = (
        "h6",
        "additional details",
        "Additional Details",
        ":",
        True,
        False,
    )
    HIGHLIGHTS = ("h3", "highlights", "Highlights", "", True, False)

    def __init__(
        self,
        markdown_heading_level: str,
        markdown_text_lower: str,
        docx_heading: str,
        separator: str = "",
        bold: bool = True,
        italic: bool = False,
    ):
        """Initialize the JobSubsection with its properties"""
        self.markdown_heading_level = markdown_heading_level
        self.markdown_text_lower = markdown_text_lower
        self.docx_heading = docx_heading
        self.separator = separator
        self.bold = bold
        self.italic = italic

    @property
    def full_heading(self):
        """Return the full heading with separator

        Returns:
            str: The complete heading text with separator
        """
        return f"{self.docx_heading}{self.separator}"

    @classmethod
    def find_by_tag_and_text(cls, tag_name: str, text: str):
        """Find a JobSubsection by tag name and text content (case insensitive)

        Args:
            tag_name (str): HTML tag name to match (e.g., 'h5', 'h6')
            text (str): Text content to match (case insensitive)

        Returns:
            JobSubsection or None: The matching subsection or None if not found
        """
        text_lower = text.lower().strip()

        # First try exact match
        for subsection in cls:
            if (
                subsection.markdown_heading_level == tag_name
                and text_lower == subsection.markdown_text_lower
            ):
                return subsection

        # If no exact match, try partial match for elements
        # Create a dictionary of keywords to check against each subsection type
        if tag_name in [
            "h5",
            "h6",
        ]:  # Allow partial matching for both h5 and h6 elements
            # Group subsections by heading level
            matching_subsections = [
                s for s in cls if s.markdown_heading_level == tag_name
            ]

            # Try partial matching against each subsection's keywords
            for subsection in matching_subsections:
                # Extract the first word of the markdown_text_lower as the key keyword
                keyword = subsection.markdown_text_lower.split()[0]
                if keyword in text_lower:
                    return subsection

        return None


class PdfConverterPaths(Enum):
    """Enum to store paths to PDF converter executables on different platforms"""

    # Windows LibreOffice paths
    LIBREOFFICE_WINDOWS = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]

    # macOS LibreOffice path
    LIBREOFFICE_MACOS = "/Applications/LibreOffice.app/Contents/MacOS/soffice"

    # Linux/Unix LibreOffice command name (assumed to be in PATH)
    LIBREOFFICE_LINUX = "libreoffice"


##############################
# Helper Classes
##############################
class ResumeSection:
    """Dynamic resume section configuration based on YAML config"""

    _sections = {}
    _section_order = []  # Preserve order from YAML
    _initialized = False

    def __init__(self, key: str, config: Dict[str, str], order_index: int):
        """Initialize a resume section from configuration

        Args:
            key: The section key (e.g., 'about', 'skills')
            config: Dictionary containing section configuration
            order_index: The position in the YAML file (0-based)
        """
        self.key = key
        self.markdown_heading = config["markdown_heading"]
        self.docx_heading = config["docx_heading"]
        self.markdown_heading_lower = self.markdown_heading.lower()
        self.add_space_before_h2 = config.get("add_space_before_h2", False)
        self.space_before_h2 = config.get("space_before_h2", None)
        self.space_after_h2 = config.get("space_after_h2", None)
        self.add_space_before_h3 = config.get("add_space_before_h3", False)
        self.space_before_h3 = config.get("space_before_h3", None)
        self.space_after_h3 = config.get("space_after_h3", None)
        self.add_space_before_h4 = config.get("add_space_before_h4", False)
        self.space_before_h4 = config.get("space_before_h4", None)
        self.space_after_h4 = config.get("space_after_h4", None)
        self.order = order_index  # Use position in YAML as order

    def matches(self, text):
        """Check if the given text matches this section's markdown_heading (case insensitive)

        Args:
            text (str): Text to compare against markdown_heading

        Returns:
            bool: True if text matches markdown_heading (case insensitive), False otherwise
        """
        return text.lower() == self.markdown_heading_lower

    @classmethod
    def init_from_config(cls, resume_sections_config: Dict[str, str]):
        """Initialize all resume sections from configuration

        Args:
            resume_sections_config: Ordered dictionary containing all section configurations
                                  (order preserved from YAML)
        """
        cls._sections = {}
        cls._section_order = []

        # Process sections in the order they appear in the YAML
        for order_index, (key, config) in enumerate(resume_sections_config.items()):
            section = cls(key, config, order_index)
            cls._sections[key.upper()] = section
            cls._section_order.append(section)

        cls._initialized = True

    @classmethod
    def get_section(cls, key: str):
        """Get a section by key

        Args:
            key: Section key (case insensitive)

        Returns:
            ResumeSection instance or None if not found
        """
        cls._check_initialized()
        return cls._sections.get(key.upper())

    @classmethod
    def get_ordered_sections(cls):
        """Get all resume sections in the order they appear in the YAML config

        Returns:
            list: List of ResumeSection instances in YAML order
        """
        cls._check_initialized()
        return cls._section_order.copy()

    @classmethod
    def all_sections(cls):
        """Get all sections as a dictionary

        Returns:
            dict: Dictionary of section_key -> ResumeSection
        """
        cls._check_initialized()
        return cls._sections.copy()

    @classmethod
    def _check_initialized(cls):
        """Check if sections have been initialized from config"""
        if not cls._initialized:
            raise RuntimeError(
                "ResumeSection has not been initialized. Call ResumeSection.init_from_config() first."
            )

    def __str__(self):
        return (
            f"ResumeSection({self.key}: {self.markdown_heading} -> {self.docx_heading})"
        )

    def __repr__(self):
        return self.__str__()


class OutputFilePath:
    """Class to handle output file path generation

    Properties:
        input_file (Path): The input file path
        output_file (Path): The output file path
        extension (str): The file extension for the output file
        interactive (bool): Whether to run in interactive mode
    """

    def __init__(
        self,
        input_file: Path,
        output_file: Path = None,
    ):
        self.input_file = input_file
        self.output_file = output_file

    def output_path(
        self, extension: str = DOCX_EXTENSION, interactive: bool = False
    ) -> Path:
        """Get the output file path

        Args:
            extension (str): The file extension for the output file
            interactive (bool): Whether to run in interactive mode

        Returns:
            The output path
        """

        default_output_name = self.input_file.with_suffix(f".{extension}").name
        default_output_file = os.path.join(DEFAULT_OUTPUT_DIR, default_output_name)
        output_path = self.output_file

        if interactive:
            # Prompt for output file
            output_prompt = f"📝 Enter the output docx filename (default: '{default_output_file}'): "
            output_path = input(output_prompt).strip()

        output_file = output_path or default_output_file

        if interactive:
            if not self.output_file:
                print(f"✅ Using default output: {default_output_file}")

        return output_file


class ConfigLoader:
    """Class for loading and accessing configuration from YAML file"""

    def __init__(
        self, config_file: Path = DEFAULT_CONFIG_FILE, print_success_msg: bool = True
    ):
        """Initialize by loading configuration from YAML file

        Args:
            config_file (Path): Path to configuration file.
                             Defaults to 'resume_config.yaml' in same directory.
            print_success_msg (bool): Whether to print success message after loading.
                                   Defaults to False.
        """
        from collections import OrderedDict

        import yaml
        from yaml.parser import ParserError

        # Default empty configuration structure
        self._config = {
            "document_defaults": {},
            "style_constants": {},
            "document_styles": {},
            "paragraph_lists": {},
            "paragraph_headings": {},
            "resume_sections": OrderedDict(),
        }

        # Try to load the YAML config file
        if os.path.exists(config_file):
            try:
                # Use a custom YAML loader that preserves order
                class OrderedLoader(yaml.SafeLoader):
                    pass

                def construct_mapping(loader, node):
                    loader.flatten_mapping(node)
                    return OrderedDict(loader.construct_pairs(node))

                OrderedLoader.add_constructor(
                    yaml.resolver.BaseResolver.DEFAULT_MAPPING_TAG, construct_mapping
                )

                with open(config_file, "r") as f:
                    yaml_config = yaml.load(f, OrderedLoader)

                if yaml_config and isinstance(yaml_config, dict):
                    # Replace resume_sections if provided (preserving order)
                    if "resume_sections" in yaml_config:
                        self._config["resume_sections"] = yaml_config["resume_sections"]

                    # Validate document styles after loading
                    if "document_styles" in yaml_config:
                        validated_styles = {}
                        for style_name, properties in yaml_config[
                            "document_styles"
                        ].items():
                            validated_styles[style_name] = _validate_style_properties(
                                properties
                            )
                        self._config["document_styles"] = validated_styles

                    # Update other sections as before...
                    for section in [
                        "document_defaults",
                        "style_constants",
                        "paragraph_lists",
                        "paragraph_headings",
                    ]:
                        if section in yaml_config:
                            if isinstance(yaml_config[section], dict) and isinstance(
                                self._config.get(section, {}), dict
                            ):
                                self._config[section].update(yaml_config[section])
                            else:
                                self._config[section] = yaml_config[section]

                    if print_success_msg:
                        print(f"✅ Config loaded from {config_file}")
            except (ParserError, Exception) as e:
                print(
                    f"❌ Error loading config file: {str(e)}, using default configuration"
                )

    @property
    def two_column_enabled(self) -> bool:
        """Check if two column layout is enabled"""
        return self._config.get("document_defaults", {}).get(
            "two_column_enabled", False
        )

    @property
    def config(self) -> dict:
        """Get the entire configuration dictionary

        Returns:
            dict: Complete configuration dictionary
        """
        return self._config

    @property
    def document_defaults(self) -> dict:
        """Get document default settings

        Returns:
            dict: Document defaults configuration
        """
        return self._config.get("document_defaults", {})

    @property
    def style_constants(self) -> dict:
        """Get style constants

        Returns:
            dict: Style constants configuration
        """
        return self._config.get("style_constants", {})

    @property
    def document_styles(self) -> dict:
        """Get document styles

        Returns:
            dict: Document styles configuration
        """
        return self._config.get("document_styles", {})

    @property
    def paragraph_headings(self) -> dict:
        """Get markdown headings configuration

        Returns:
            dict: Markdown headings configuration
        """
        return self._config.get("paragraph_headings", {})

    @property
    def resume_sections(self) -> dict:
        """Get resume sections configuration (order preserved)

        Returns:
            dict: Resume sections configuration in YAML order
        """
        return self._config.get("resume_sections", {})

    def get_style_constant(self, key: str, default=None):
        """Get a specific style constant value

        Args:
            key (str): Style constant key to retrieve
            default: Default value if key is not found

        Returns:
            Value for the requested style constant or default if not found
        """
        return self.style_constants.get(key, default)


class ConfigHelper:
    """Static helper class for accessing configuration values globally"""

    # Class-level storage for configuration
    _config = None
    _initialized = False

    @classmethod
    def init(cls, config: Dict[str, dict]) -> None:
        """Initialize configuration store

        Args:
            config (dict): Complete configuration dictionary
        """
        cls._config = config
        cls._initialized = True

    @classmethod
    def get_document_defaults(cls) -> dict:
        """Get document default settings

        Returns:
            dict: Document defaults configuration
        """
        cls._check_initialized()
        return cls._config.get("document_defaults", {})

    @classmethod
    def get_style_constants(cls) -> dict:
        """Get style constants

        Returns:
            dict: Style constants configuration
        """
        cls._check_initialized()
        return cls._config.get("style_constants", {})

    @classmethod
    def get_document_styles(cls) -> dict:
        """Get document styles

        Returns:
            dict: Document styles configuration
        """
        cls._check_initialized()
        return cls._config.get("document_styles", {})

    @classmethod
    def get_style_constant(cls, key: str, default=None):
        """Get a specific style constant value

        Args:
            key (str): Style constant key to retrieve
            default: Default value if key is not found

        Returns:
            Value for the requested style constant or default if not found
        """
        cls._check_initialized()
        return cls._config.get("style_constants", {}).get(key, default)

    @classmethod
    def get_paragraph_list_option(cls, list_type: str, option_name: str, default=None):
        """Get a paragraph list option from config

        Args:
            list_type: The list type ("ul" or "ol")
            option_name: The option name to retrieve
            default: Default value if not found

        Returns:
            The option value or default
        """
        cls._check_initialized()
        list_config = cls._config.get("paragraph_lists", {})
        type_config = list_config.get(list_type, {})
        return type_config.get(option_name, default)

    @classmethod
    def _check_initialized(cls) -> None:
        """Check if the configuration has been initialized

        Raises:
            RuntimeError: If the configuration has not been initialized
        """
        if not cls._initialized:
            raise RuntimeError(
                "ConfigHelper has not been initialized. Call ConfigHelper.init() first."
            )


class HeadingsHelper:
    """Helper class for managing markdown headings based on configuration"""

    # Class-level storage for the heading configuration
    _heading_map = {}
    _paragraph_style_headings = {}
    _initialized = False

    @classmethod
    def init(
        cls, config: Dict[str, dict], paragraph_style_headings: Dict[str, bool] = None
    ) -> None:
        """Initialize the heading map from configuration

        Args:
            config (dict): Configuration dictionary with paragraph_headings
        """
        # Simply assign the paragraph_headings from config
        cls._heading_map = config["paragraph_headings"]
        cls._paragraph_style_headings = paragraph_style_headings or {}
        cls._initialized = True

    @classmethod
    def should_use_paragraph_style(cls, tag_name: str) -> bool:
        """Check if a heading tag should use paragraph style instead of heading style

        Args:
            tag_name (str): HTML tag name (e.g., 'h1', 'h2', etc.)

        Returns:
            bool: True if tag should use paragraph style, False otherwise
        """
        cls._check_initialized()
        return cls._paragraph_style_headings.get(tag_name, False)

    @classmethod
    def any_heading_uses_paragraph_style(cls) -> bool:
        """Check if any heading level uses paragraph style

        Returns:
            bool: True if any heading level has paragraph styling enabled, False otherwise
        """
        cls._check_initialized()
        return bool(cls._paragraph_style_headings)

    @classmethod
    def get_level_for_tag(cls, tag_name: str) -> int | None:
        """Get the Word document heading level for a given tag name

        Args:
            tag_name (str): HTML tag name (e.g., 'h1', 'h2', etc.)

        Returns:
            int or None: The corresponding Word document heading level or None if not found
        """
        cls._check_initialized()
        if tag_name.lower() in cls._heading_map:
            return cls._heading_map[tag_name.lower()]["level"]
        return None

    @classmethod
    def get_font_size_for_level(cls, heading_level: int, default_size: int = 11) -> int:
        """Get the font size for a given heading level

        Args:
            heading_level (int): The heading level (0-5)
            default_size (int): The default font size to return if not found

        Returns:
            int: The font size in points
        """
        cls._check_initialized()
        # Find the heading tag that corresponds to this level
        for tag, props in cls._heading_map.items():
            if props["level"] == heading_level:
                return props["paragraph_heading_size"]
        return default_size

    @classmethod
    def get_font_size_for_tag(cls, tag_name: str, default_size: int = 11) -> int:
        """Get the font size for a given tag name directly

        Args:
            tag_name (str): HTML tag name (e.g., 'h1', 'h2', etc.)
            default_size (int): The default font size to return if not found

        Returns:
            int: The font size in points for the tag
        """
        cls._check_initialized()
        tag_name = tag_name.lower()
        if tag_name in cls._heading_map:
            return cls._heading_map[tag_name]["paragraph_heading_size"]
        return default_size

    @classmethod
    def _check_initialized(cls) -> None:
        """Check if the configuration has been initialized

        Raises:
            RuntimeError: If the configuration has not been initialized
        """
        if not cls._initialized:
            raise RuntimeError(
                "ConfigHelper has not been initialized. Call ConfigHelper.init() first."
            )


##############################
# Main Processors
##############################
def create_ats_resume(
    md_file: Path,
    output_file: Path,
    config_loader: ConfigLoader,
    paragraph_style_headings: Dict[str, bool] = None,
) -> Path:
    """Convert markdown resume to ATS-friendly Word document

    Args:
        md_file (Path): Path to the markdown resume file
        output_file (str): Path where the output Word document will be saved
        config_loader (ConfigLoader, optional): ConfigLoader instance with configuration.
                                               If None, creates with default config file.
        paragraph_style_headings (dict, optional): Dictionary mapping heading tags
                                                 to boolean values.

    Returns:
        Path: Path to the created document
    """
    # Default to using heading styles for all if not specified
    if paragraph_style_headings is None:
        paragraph_style_headings = {}

    # Initialize ResumeSection from config (order preserved from YAML)
    ResumeSection.init_from_config(config_loader.resume_sections)

    # Initialize ConfigHelper and HeadingsHelper with config
    ConfigHelper.init(config_loader.config)
    HeadingsHelper.init(config_loader.config, paragraph_style_headings)

    # Access config sections directly through properties
    doc_defaults = config_loader.document_defaults
    style_constants = config_loader.style_constants
    document_styles = config_loader.document_styles

    # Read markdown file
    with open(md_file, "r") as file:
        md_content = file.read()

    # Convert markdown to HTML for easier parsing
    html = markdown.markdown(md_content)
    soup = BeautifulSoup(html, "html.parser")

    # Create document with standard margins
    document = DOCX_Document()

    # Apply styles from configuration
    _apply_document_styles(document, document_styles, style_constants)

    for section in document.sections:
        section.page_width = Inches(doc_defaults["page_width"])
        section.page_height = Inches(doc_defaults["page_height"])
        section.top_margin = Inches(doc_defaults["margin_top"])
        section.bottom_margin = Inches(doc_defaults["margin_bottom"])
        section.left_margin = Inches(doc_defaults["margin_left"])
        section.right_margin = Inches(doc_defaults["margin_right"])

    contact_section = ResumeSection.get_section("CONTACT")
    about_section = ResumeSection.get_section("ABOUT")
    skills_section = ResumeSection.get_section("SKILLS")
    certifications_section = ResumeSection.get_section("CERTIFICATIONS")
    education_section = ResumeSection.get_section("EDUCATION")
    experience_section = ResumeSection.get_section("EXPERIENCE")
    projects_section = ResumeSection.get_section("PROJECTS")

    # Check if two-column layout is enabled
    if config_loader.two_column_enabled:
        # Process header with name and tagline
        process_header_section(document, soup)

        # Create the two-column layout structure
        about_content_cell, contact_info_cell, sidebar_cell, main_cell = (
            _create_two_column_layout(document)
        )

        # Process contact info in the contact cell (right after header, before about section)
        contact_section = ResumeSection.get_section("CONTACT")
        if contact_section:
            _process_contact_info_horizontal(contact_info_cell, soup)

        about_section = ResumeSection.get_section("ABOUT")
        if about_section:
            # Override space before to be minimal
            about_section.space_before_h2 = 0
            process_about_section(about_content_cell, soup)

        # Process sections for the sidebar (Top Skills, Certifications, Education)
        sidebar_sections = []
        first_sidebar_section = True

        # Add Top Skills to sidebar
        if skills_section:
            process_skills_section(sidebar_cell, soup)
            sidebar_sections.append(skills_section.docx_heading)
            first_sidebar_section = False

            # Add horizontal line after skills if there are more sections coming
            if certifications_section or education_section:
                _add_horizontal_line_to_cell(sidebar_cell, is_sidebar=True)

        # Add Certifications to sidebar
        if certifications_section:
            # If this isn't the first section and no line added yet, add one now
            if not first_sidebar_section and (
                sidebar_sections[-1] != skills_section.docx_heading
            ):
                _add_horizontal_line_to_cell(sidebar_cell, is_sidebar=True)

            process_certifications_section(sidebar_cell, soup)
            sidebar_sections.append(certifications_section.docx_heading)
            first_sidebar_section = False

            # Add horizontal line after certifications if education is coming
            if education_section:
                _add_horizontal_line_to_cell(sidebar_cell, is_sidebar=True)

        # Add Education to sidebar
        if education_section:
            # If this isn't the first section and no line added yet, add one now
            if not first_sidebar_section and (
                sidebar_sections[-1] != certifications_section.docx_heading
            ):
                _add_horizontal_line_to_cell(sidebar_cell, is_sidebar=True)

            process_education_section(sidebar_cell, soup)
            sidebar_sections.append(education_section.docx_heading)

        # Track main sections for horizontal lines
        main_sections = []
        first_main_section = True

        # Process Experience section
        if experience_section:
            # Add horizontal line before first section
            # _add_horizontal_line_to_cell(main_cell, is_sidebar=False)

            process_experience_section(main_cell, soup)
            main_sections.append(experience_section.docx_heading)
            first_main_section = False

        # Process Projects section
        if projects_section:
            # Add horizontal line before projects section
            if not first_main_section:
                _add_horizontal_line_to_cell(main_cell, is_sidebar=False)

            process_projects_section(main_cell, soup)
            main_sections.append(projects_section.docx_heading)

        # print("Applying styles to all cells...")
        # FIXME: This overrides any custom styles applied in the section processors, which is giving a particular problem for the sidebar horizontal line spacing
        for cell in [about_content_cell, contact_info_cell, sidebar_cell, main_cell]:
            _apply_styles_to_cell_content(cell, document_styles)
        # print("Styles applied.")
    else:
        # Define processor mapping using dynamic section access
        section_processor_map = {}

        # Get sections dynamically
        if about_section:
            section_processor_map[about_section] = [
                (process_header_section, True),  # Header always required
                (
                    lambda doc, soup: process_about_section(document=doc, soup=soup),
                    False,
                ),
            ]

        if skills_section:
            section_processor_map[skills_section] = [
                (lambda doc, soup: process_skills_section(doc, soup), False),
            ]

        if experience_section:
            section_processor_map[experience_section] = [
                (lambda doc, soup: process_experience_section(doc, soup), False),
            ]

        if projects_section:
            section_processor_map[projects_section] = [
                (lambda doc, soup: process_projects_section(doc, soup), False),
            ]

        if certifications_section:
            section_processor_map[certifications_section] = [
                (lambda doc, soup: process_certifications_section(doc, soup), False),
            ]

        if education_section:
            section_processor_map[education_section] = [
                (lambda doc, soup: process_education_section(doc, soup), False),
            ]

        if contact_section:
            section_processor_map[contact_section] = [
                (lambda doc, soup: process_contact_section(doc, soup), False),
            ]

        # Build section processors in YAML order
        section_processors = []
        for section_type in ResumeSection.get_ordered_sections():
            if section_type in section_processor_map:
                for processor, required in section_processor_map[section_type]:
                    section_processors.append((section_type, processor, required))

        # Process each section with error handling
        for section_type, processor, required in section_processors:
            try:
                processor(document, soup)
            except Exception as e:
                if required:
                    raise  # Re-raise for required sections
                else:
                    print(
                        f"⚠️  Warning: Could not process {section_type.docx_heading}: {str(e)}"
                    )

    # Add page numbers if enabled
    _add_configured_page_numbers(document)

    # Save the document
    document.save(output_file)

    return Path(output_file)


def convert_to_pdf(docx_file: Path) -> Path | None:
    """Convert a DOCX file to PDF using available converters

    Args:
        docx_file (Path): Path to the input DOCX file

    Returns:
        Path: Path to the created PDF file, or None if conversion failed
    """
    pdf_file = docx_file.with_suffix(f".{PDF_EXTENSION}")

    # Try multiple conversion methods
    methods = [
        _convert_with_docx2pdf,
        _convert_with_libreoffice,
        _convert_with_win32com,
    ]

    for method in methods:
        try:
            if method(docx_file, pdf_file):
                return Path(pdf_file)
        except Exception as e:
            print(f"Could not convert using {method.__name__}: {str(e)}")

    print("❌ PDF conversion failed. Please install one of the following:")
    print("   - docx2pdf (pip install docx2pdf)")
    print("   - LibreOffice (https://www.libreoffice.org/)")
    print("   - Microsoft Word (Windows only)")
    return None


##############################
# Section Processors
##############################
def process_header_section(
    document: DOCX_Document,
    soup: BeautifulSoup,
) -> None:
    """Process the header (name and tagline) section

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content

    Returns:
        None
    """
    # Get header image settings
    header_defaults = ConfigHelper.get_document_defaults()
    header_image_enabled = header_defaults.get("header_image", {}).get("enabled", False)
    header_alignment = header_defaults.get("header_alignment", "center")

    # Always extract and add the main title (name) directly to the document
    name = soup.find("h1").text
    title_para = document.add_paragraph(style="Title")
    # Always center the title regardless of header_alignment setting
    _paragraph_alignment(title_para, "center")
    title_para.add_run(name)

    # Extract the tagline (first paragraph after h1)
    first_p = soup.find("h1").find_next_sibling()
    has_tagline = first_p and first_p.name == "p"

    # Create a table for the image and tagline only if image is enabled
    if header_image_enabled:

        # Create a 1x2 table for the header (image on left, tagline on right)
        header_table = _create_table(document, 1, 2)

        # Get image settings
        header_image = header_defaults.get("header_image", {})

        img_url = header_image.get("url", "")
        img_width = Inches(header_image.get("width_inches", 1.0))
        img_height = Inches(header_image.get("height_inches", 1.0))
        img_border_color = header_image.get("border_color", "000000")
        img_border_width = header_image.get("border_width", 1)
        margin_right = Inches(header_image.get("margin_right", 0.5))

        # Calculate available width and set column widths
        page_width = header_defaults.get("page_width", 8.5)
        margin_left = header_defaults.get("margin_left", 0.5)
        margin_right_page = header_defaults.get("margin_right", 0.5)
        available_width = page_width - margin_left - margin_right_page

        # Set first column to exact image width plus margin
        first_col_width = img_width + margin_right
        second_col_width = Inches(available_width) - first_col_width

        header_table.columns[0].width = first_col_width
        header_table.columns[1].width = second_col_width

        # Add image to first cell
        img_cell = header_table.cell(0, 0)
        img_align_vertical = header_image.get("vertical_alignment", "center")
        _cell_vertical_alignment(img_cell, img_align_vertical)

        # Clear any existing paragraph
        if img_cell.paragraphs:
            p = img_cell.paragraphs[0]
            p._element.getparent().remove(p._element)
            p._p = None
            p._element = None

        img_para = img_cell.add_paragraph()

        # Add image
        _url_image(
            img_para,
            img_url,
            img_width,
            img_height,
            img_border_color,
            img_border_width,
        )

        # Use second cell for tagline
        text_cell = header_table.cell(0, 1)
        header_align_vertical = header_defaults.get("vertical_alignment", "center")
        _cell_vertical_alignment(text_cell, header_align_vertical)

        if text_cell.paragraphs:
            p = text_cell.paragraphs[0]
            p._element.getparent().remove(p._element)
            p._p = None
            p._element = None

        # Process tagline in the right cell if it exists
        if has_tagline:
            _process_tagline_and_specialty(text_cell, first_p, header_alignment)

        # Add horizontal line to the tagline cell instead of to the document
        # This centers it below just the tagline portion
        _add_horizontal_line_to_cell(text_cell)
    else:
        # If no image, just process the tagline directly in the document
        if has_tagline:
            _process_tagline_and_specialty(document, first_p, header_alignment)

        # Add horizontal line to the full document when no image
        _add_horizontal_line_simple(document)


def process_about_section(
    document: DOCX_Document,
    soup: BeautifulSoup,
) -> None:
    """Process the About section

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content

    Returns:
        None
    """
    about_section = ResumeSection.get_section("ABOUT")
    add_space_before_h2 = about_section.add_space_before_h2
    space_before_h2 = about_section.space_before_h2 if add_space_before_h2 else None
    space_after_h2 = about_section.space_after_h2

    section_h2 = _prepare_section(
        document,
        soup,
        about_section,
        space_before=space_before_h2,
        space_after=space_after_h2,
    )

    add_space_before_h3 = about_section.add_space_before_h3
    space_before_h3 = about_section.space_before_h3 if add_space_before_h3 else None
    space_after_h3 = about_section.space_after_h3

    if not section_h2:
        return  # Gracefully exit if section doesn't exist

    current_element = section_h2.find_next_sibling()

    # Initialize processed_elements set if we're in the About section
    processed_elements = set()

    # Add all paragraphs until next h2
    while current_element and current_element.name != "h2":
        # Skip if already processed
        if current_element in processed_elements:
            current_element = current_element.find_next_sibling()
            continue

        # Check if this is a paragraph with strong element containing highlights
        highlights_subsection = None
        if (
            current_element.name == JobSubsection.HIGHLIGHTS.markdown_heading_level
            and current_element.text.strip().lower()
            == JobSubsection.HIGHLIGHTS.markdown_text_lower
        ):
            highlights_subsection = JobSubsection.HIGHLIGHTS

        # Handle regular paragraph
        if not highlights_subsection and current_element.name == "p":
            para = document.add_paragraph()

            # Process all elements of the paragraph to preserve formatting
            for child in current_element.children:
                # Check if this is a strong/bold element
                if getattr(child, "name", None) == "strong":
                    run = para.add_run(child.text)
                    run.bold = True
                # Check if this is an em/italic element
                elif getattr(child, "name", None) == "em":
                    run = para.add_run(child.text)
                    run.italic = True
                # Check if this is a link/anchor element
                elif getattr(child, "name", None) == "a" and child.get("href"):
                    _add_hyperlink(para, child.text, child.get("href"))
                # Otherwise, just add the text as-is
                else:
                    if child.string:
                        _process_text_for_hyperlinks(para, child.string)

            processed_elements.add(current_element)

        # Handle highlights subsection found in a paragraph or heading
        elif highlights_subsection:
            heading_level = (
                HeadingsHelper.get_level_for_tag(current_element.name)
                if current_element.name.startswith("h")
                else None
            )
            use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
                current_element.name
            )
            _add_heading_or_paragraph(
                document,
                highlights_subsection.full_heading,
                heading_level,
                use_paragraph_style=use_paragraph_style,
                bold=highlights_subsection.bold,
                italic=highlights_subsection.italic,
                space_before=space_before_h3,
                space_after=space_after_h3,
            )
            processed_elements.add(current_element)

        # Handle bullet list
        elif current_element.name == "ul":
            _add_bullet_list(document, current_element)
            processed_elements.add(current_element)

        current_element = current_element.find_next_sibling()


def process_skills_section(
    document: DOCX_Document,
    soup: BeautifulSoup,
) -> None:
    """Process the Skills section

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content

    Returns:
        None
    """
    skills_section = ResumeSection.get_section("SKILLS")
    section_h2 = _prepare_section(document, soup, skills_section)

    if not section_h2:
        return  # Gracefully exit if section doesn't exist

    current_element = section_h2.find_next_sibling()

    if current_element and current_element.name == "p":
        _process_horizontal_skills_list(
            document, current_element.text, is_top_skills=True
        )


def process_experience_section(
    document: DOCX_Document,
    soup: BeautifulSoup,
) -> None:
    """Process the Experience section

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content

    Returns:
        None
    """
    experience_section = ResumeSection.get_section("EXPERIENCE")
    section_h2 = _prepare_section(
        document,
        soup,
        experience_section,
        space_after=experience_section.space_after_h2,
    )

    if not section_h2:
        return  # Gracefully exit if section doesn't exist

    # Find all job entries (h3 headings under Experience)
    current_element = section_h2.find_next_sibling()
    # Use a set of element IDs instead of element objects
    processed_element_ids = set()
    processed_elements = set()
    add_space_before_h3 = experience_section.add_space_before_h3
    space_before_h3 = (
        experience_section.space_before_h3 if add_space_before_h3 else None
    )
    space_after_h3 = experience_section.space_after_h3

    add_space_before_h4 = experience_section.add_space_before_h4
    space_before_h4 = (
        experience_section.space_before_h4 if add_space_before_h4 else None
    )
    space_after_h4 = experience_section.space_after_h4

    # Track if this is the first job
    first_job = True

    while current_element and current_element.name != "h2":
        # Get unique ID for this element
        element_id = id(current_element)

        # Skip if already processed (except h3 elements)
        if element_id in processed_element_ids and current_element.name != "h3":
            current_element = current_element.find_next_sibling()
            continue

        # Process based on element type
        if current_element.name == "h3":
            # Process job entry and mark as processed
            processed_elements = _process_job_entry(
                document,
                current_element,
                set(),
                space_before_h3,
                space_after_h3,
                is_first_job=first_job,
            )
            processed_element_ids.add(element_id)

            if first_job:
                first_job = False

        elif current_element.name == "h4" and element_id not in processed_element_ids:
            # Process position and mark as processed
            processed_elements = _process_position(
                document, current_element, set(), space_before_h4, space_after_h4
            )
            processed_element_ids.add(element_id)

        elif (
            current_element.name in ["h5", "h6"]
            and element_id not in processed_element_ids
        ):
            # Find matching subsection type
            subsection = JobSubsection.find_by_tag_and_text(
                current_element.name, current_element.text
            )

            if subsection:
                heading_level = HeadingsHelper.get_level_for_tag(current_element.name)

                # PROJECT/CLIENT requires special handling with its own function
                if (
                    subsection == JobSubsection.PROJECT_CLIENT
                    and current_element not in processed_elements
                ):
                    project_processed = _process_project_section(
                        document,
                        current_element,
                        processed_elements,
                    )

                    # Add all elements processed by the project handler to our tracking
                    processed_elements.update(project_processed)

                    # Also update element IDs
                    for element in project_processed:
                        processed_element_ids.add(id(element))

                # Generic subsection processing for SUMMARY, INTERNAL, etc.
                elif (
                    subsection in [JobSubsection.SUMMARY, JobSubsection.INTERNAL]
                    and current_element not in processed_elements
                ):
                    project_processed = _process_subsection(
                        document,
                        current_element,
                        subsection,
                        heading_level,
                        processed_elements,
                    )

                    # Update tracking
                    processed_elements.update(project_processed)
                    for element in project_processed:
                        processed_element_ids.add(id(element))

                # KEY_SKILLS subsection
                elif subsection == JobSubsection.KEY_SKILLS:
                    use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
                        current_element.name
                    )
                    key_skills_heading_line_spacing = ConfigHelper.get_style_constant(
                        "key_skills_heading_line_spacing", None
                    )
                    skills_heading = _add_heading_or_paragraph(
                        document,
                        subsection.full_heading,
                        heading_level,
                        use_paragraph_style=use_paragraph_style,
                        bold=subsection.bold,
                        italic=subsection.italic,
                    )

                    # Set line spacing for the heading
                    _apply_paragraph_format_properties(
                        skills_heading.paragraph_format,
                        {"line_spacing": key_skills_heading_line_spacing},
                    )

                    # Get skills from next element
                    next_element = current_element.find_next_sibling()
                    if next_element and next_element.name == "p":
                        skills_para = _process_horizontal_skills_list(
                            document, next_element.text, is_top_skills=False
                        )
                        processed_elements.add(next_element)

                    # Determine if we need to add a blank line after Key Skills
                    # We'll add a blank line if:
                    # 1. There are no more headings (end of section)
                    # 2. The next heading is an h3 (new job)
                    # 3. The next heading is an h4 (new role within same company)
                    # We won't add a blank line if there's an h5 or h6 after Key Skills
                    looking_ahead = current_element
                    next_heading = None

                    # Look for the next heading element
                    while looking_ahead and not next_heading:
                        looking_ahead = looking_ahead.find_next_sibling()
                        if looking_ahead and looking_ahead.name in [
                            "h3",
                            "h4",
                            "h5",
                            "h6",
                        ]:
                            next_heading = looking_ahead

                    # Add space if this is the last role or before a new role (most likely h4)
                    # TODO: can this be removed?
                    # if not next_heading or next_heading.name in ["h4"]:
                    #     _add_space_paragraph(document, 8)

                # RESPONSIBILITIES subsection (standalone)
                elif (
                    subsection == JobSubsection.RESPONSIBILITIES
                    and current_element not in processed_elements
                ):
                    use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
                        current_element.name
                    )
                    _add_heading_or_paragraph(
                        document,
                        subsection.full_heading,
                        heading_level,
                        use_paragraph_style=use_paragraph_style,
                        bold=subsection.bold,
                        italic=subsection.italic,
                    )

                    # Get content
                    next_element = current_element.find_next_sibling()
                    if next_element:
                        if next_element.name == "p":
                            resp_para = document.add_paragraph()
                            _process_text_for_hyperlinks(resp_para, next_element.text)
                            processed_elements.add(next_element)
                        elif next_element.name == "ul":
                            # Process bullet list
                            _add_bullet_list(document, next_element)
                            processed_elements.add(next_element)

                # ADDITIONAL_DETAILS subsection (standalone)
                elif (
                    subsection == JobSubsection.ADDITIONAL_DETAILS
                    and current_element not in processed_elements
                ):
                    use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
                        current_element.name
                    )
                    _add_heading_or_paragraph(
                        document,
                        subsection.full_heading,
                        heading_level,
                        use_paragraph_style=use_paragraph_style,
                        bold=subsection.bold,
                        italic=subsection.italic,
                    )

                    # Get content (next element might be list items)
                    next_element = current_element.find_next_sibling()
                    if next_element and next_element.name == "ul":
                        _add_bullet_list(document, next_element)
                        processed_elements.add(next_element)

        # Standalone bullet points
        elif current_element.name == "ul" and current_element not in processed_elements:
            # Process bullet list
            _add_bullet_list(document, current_element)
            processed_element_ids.add(element_id)

        current_element = current_element.find_next_sibling()


def process_education_section(
    document: DOCX_Document,
    soup: BeautifulSoup,
) -> None:
    """Process the Education section

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content

    Returns:
        None
    """
    education_section = ResumeSection.get_section("EDUCATION")
    section_h2 = _prepare_section(
        document,
        soup,
        education_section,
    )

    if not section_h2:
        return  # Gracefully exit if section doesn't exist

    _process_simple_section(
        document,
        section_h2,
        add_space=education_section.add_space_before_h3,
    )


def process_certifications_section(
    document: DOCX_Document,
    soup: BeautifulSoup,
) -> None:
    """Process the Certifications section

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content

    Returns:
        None
    """
    certifications_section = ResumeSection.get_section("CERTIFICATIONS")
    section_h2 = _prepare_section(
        document,
        soup,
        certifications_section,
    )

    if not section_h2:
        return  # Gracefully exit if section doesn't exist

    add_space_before_h3 = certifications_section.add_space_before_h3
    space_before_h3 = (
        certifications_section.space_before_h3 if add_space_before_h3 else None
    )
    space_after_h3 = certifications_section.space_after_h3

    add_space_before_h4 = certifications_section.add_space_before_h4
    space_before_h4 = (
        certifications_section.space_before_h4 if add_space_before_h4 else None
    )
    space_after_h4 = certifications_section.space_after_h4

    _process_projects_or_certifications(
        document,
        section_h2,
        space_before_h3=space_before_h3,
        space_after_h3=space_after_h3,
        space_before_h4=space_before_h4,
        space_after_h4=space_after_h4,
    )


def process_projects_section(
    document: DOCX_Document,
    soup: BeautifulSoup,
) -> None:
    """Process the Special Projects section

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content

    Returns:
        None
    """
    projects_section = ResumeSection.get_section("PROJECTS")
    section_h2 = _prepare_section(
        document,
        soup,
        projects_section,
    )

    if not section_h2:
        return  # Gracefully exit if section doesn't exist

    add_space_before_h3 = projects_section.add_space_before_h3
    space_before_h3 = projects_section.space_before_h3 if add_space_before_h3 else None
    space_after_h3 = projects_section.space_after_h3
    space_after_h4 = projects_section.space_after_h4

    _process_projects_or_certifications(
        document,
        section_h2,
        space_before_h3=space_before_h3,
        space_after_h3=space_after_h3,
        space_after_h4=space_after_h4,
    )


def process_contact_section(
    document: DOCX_Document,
    soup: BeautifulSoup,
) -> None:
    """Process the Contact section

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content

    Returns:
        None
    """
    contact_section = ResumeSection.get_section("CONTACT")
    section_h2 = _prepare_section(
        document,
        soup,
        contact_section,
    )

    if not section_h2:
        return  # Gracefully exit if section doesn't exist

    _process_simple_section(
        document,
        section_h2,
        add_space=contact_section.add_space_before_h3,
    )


##############################
# Two Column Helpers
##############################
def _create_table(
    document: DOCX_Document,
    rows: int = 1,
    cols: int = 1,
    style: str = "Table Grid",
    allow_autofit: bool = False,
    autofit: bool = False,
) -> DOCX_Table:
    """Create a table in the document with specified rows, columns, and style

    Args:
        document: The Word document object
        rows: Number of rows in the table
        cols: Number of columns in the table
        style: Style name for the table

    Returns:
        docx.table.Table: The created table object
    """
    table = document.add_table(rows=rows, cols=cols)
    table.allow_autofit = allow_autofit
    table.autofit = autofit

    table.style = style
    for cell in table._cells:
        cell_xml = cell._tc
        tcPr = cell_xml.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for border in ["top", "left", "bottom", "right"]:
            border_elem = OxmlElement(f"w:{border}")
            border_elem.set(qn("w:val"), "nil")
            tcBorders.append(border_elem)
        tcPr.append(tcBorders)

    return table


def _process_contact_info_horizontal(cell: DOCX_Cell, soup: BeautifulSoup) -> None:
    """Process contact info in a horizontal ribbon format with gray background

    Args:
        cell: The table cell to place contact info in
        soup: BeautifulSoup object containing the HTML content
    """

    contact_ribbon_styles = ConfigHelper.get_style_constant("contact_ribbon", {})
    item_separator = contact_ribbon_styles.get("item_separator", "  |  ")
    item_separator_color = contact_ribbon_styles.get("item_separator_color", "000000")
    item_separator_font_style = contact_ribbon_styles.get(
        "item_separator_font_style", "normal"
    )

    RGBColor.from_string(item_separator_color)

    # Set gray background for the cell
    _apply_table_cell_fill_and_border_styles(cell, contact_ribbon_styles)

    # Find the contact section
    contact_section = soup.find("h2", string=lambda text: text.lower() == "contact")
    if not contact_section:
        return

    # Create a paragraph for horizontal layout
    para = cell.add_paragraph()
    _paragraph_alignment(para, "center")

    # Find all contact entries (paragraphs after h2 until next h2)
    current_element = contact_section.find_next_sibling()
    contact_items = []

    while current_element and current_element.name != "h2":
        if current_element.name == "p":
            contact_items.append(current_element)
        current_element = current_element.find_next_sibling()

    # Add contact items horizontally with separators
    for i, item in enumerate(contact_items):
        # Add separator between items (except first)
        if i > 0:
            separator = para.add_run(item_separator)
            _apply_font_properties(
                separator.font,
                {
                    "bold": item_separator_font_style == "bold",
                    "italic": item_separator_font_style == "italic",
                    "color": item_separator_color,
                },
            )

        # Check if this item has a strong tag (label)
        strong_tag = item.find("strong")
        if strong_tag:
            # Add the label in bold
            label_run = para.add_run(strong_tag.text.strip())
            _apply_font_properties(
                label_run.font,
                {
                    "bold": True,
                },
            )

            # Add the value (rest of the text)
            value_text = item.text.replace(strong_tag.text, "").strip()
            if value_text:
                para.add_run(f" {value_text}")
        else:
            # Process the entire text
            _process_text_for_hyperlinks(para, item.text.strip())


def _apply_styles_to_cell_content(
    cell: DOCX_Cell, document_styles: Dict[str, dict]
) -> None:
    """Apply document styles directly to table cell content following Stack Overflow approach

    Args:
        cell: The table cell containing content
        document_styles: Dictionary of document styles from config
    """

    # StackOverflow approach: Create a direct map of style properties to apply
    style_properties = {
        "Normal": document_styles.get("Normal", {}),
        "Heading 1": document_styles.get("Heading 1", {}),
        "Heading 2": document_styles.get("Heading 2", {}),
        "Heading 3": document_styles.get("Heading 3", {}),
        "Heading 4": document_styles.get("Heading 4", {}),
        "Heading 5": document_styles.get("Heading 5", {}),
        "Hyperlink": document_styles.get("Hyperlink", {}),
    }

    # Define section heading patterns for better heading detection - make this more specific
    heading_patterns = {
        r"^(PROFESSIONAL\s+SUMMARY|ABOUT|SUMMARY)$": 2,
        r"^(PROFESSIONAL\s+EXPERIENCE|WORK\s+EXPERIENCE|EXPERIENCE)$": 2,
        r"^(TOP\s+SKILLS|SKILLS|EXPERTISE)$": 2,
        r"^(EDUCATION|ACADEMIC)$": 2,
        r"^(CONTACT|CONTACT\s+INFORMATION)$": 2,
        r"^(PROJECTS|SPECIAL\s+PROJECTS)$": 2,
        r"^(LICENSES|CERTIFICATIONS|LICENSES\s+&\s+CERTIFICATIONS)$": 2,
    }

    # Process each paragraph in the cell
    for i, paragraph in enumerate(cell.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        # Start with Normal style as base
        normal = style_properties["Normal"]

        # Detect heading level using various methods
        heading_level = None

        # Method 1: Check if all runs are bold (common heading indicator)
        is_bold = all(run.bold for run in paragraph.runs if run.text.strip())

        # Method 2: Pattern match for section headings - this is the key fix!
        for pattern, level in heading_patterns.items():
            if re.match(pattern, text, re.IGNORECASE):
                heading_level = 1
                break

        # Method 3: Check font size of first run if available
        if not heading_level and paragraph.runs and is_bold:
            first_run = paragraph.runs[0]
            if (
                hasattr(first_run, "font")
                and hasattr(first_run.font, "size")
                and first_run.font.size
            ):
                font_size = first_run.font.size.pt
                if font_size >= 16:
                    heading_level = 1  # H1
                elif font_size >= 14:
                    heading_level = 2  # H2 (main section headings)
                elif font_size >= 12:
                    heading_level = 3  # H3 (company names)
                elif font_size >= 11:
                    heading_level = 4  # H4 (position titles)

        # Method 4: Text characteristics (length, all caps, position)
        if not heading_level and is_bold:
            if text.isupper() and len(text) < 40:
                heading_level = 2  # H2 (section headings are often ALL CAPS)
            elif len(text) < 35:
                heading_level = 3  # H3 (company names tend to be shorter)
            elif text.endswith(":"):
                heading_level = 5  # H5 (subsection headings often end with colon)

        # Get the appropriate style based on detection
        style_name = f"Heading {heading_level}" if heading_level else "Normal"
        style = style_properties.get(style_name, normal)

        # Apply style to all runs in paragraph (direct approach from Stack Overflow)
        for run in paragraph.runs:
            # Bold (preserve for headings)
            if heading_level or "bold" in style:
                style["bold"] = style.get("bold", heading_level is not None)

            _apply_font_properties(run.font, style)

        _apply_paragraph_format_properties(paragraph.paragraph_format, style)

    # Special handling for hyperlinks
    hyperlink_style = style_properties.get("Hyperlink", {})
    if hyperlink_style:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                # Check for hyperlink indicators
                if run.underline or (
                    "http" in run.text or "www." in run.text or "@" in run.text
                ):
                    _apply_font_properties(run.font, hyperlink_style)


def _cell_margins(
    cell: DOCX_Cell,
    top: int = 3,
    bottom: int = 3,
    left: int = 3,
    right: int = 3,
    twips: bool = False,
) -> None:
    """Set margins for a table cell

    Args:
        cell: The table cell to set margins for
        top: Top margin in points (default 3)
        bottom: Bottom margin in points (default 3)
        left: Left margin in points (default 3)
        right: Right margin in points (default 3)

    Returns:
        The modified cell with margins set
    """
    # Get the cell's XML element
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")

    # Add each margin side
    # Convert points to twips
    side_map = {
        "top": top * 20 if not twips else top,
        "bottom": bottom * 20 if not twips else bottom,
        "start": left * 20 if not twips else left,
        "end": right * 20 if not twips else right,
    }

    for side, width in side_map.items():
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(width))  # Value is in twips
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)

    tcPr.append(tcMar)


def _cell_vertical_alignment(
    cell: DOCX_Cell, v_align: str = "top", use_xml: bool = False
) -> None:
    """Set vertical alignment for a table cell

    Args:
        cell: The table cell to set alignment for
        v_align: Vertical alignment ('top', 'center', 'bottom')

    Returns:
        The modified cell with vertical alignment set
    """
    vertical_alignment = v_align.lower()

    # Set vertical alignment
    if use_xml:
        # Get the cell's XML element
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        tcVAlign = OxmlElement("w:vAlign")
        tcVAlign.set(qn("w:val"), vertical_alignment)
        tcPr.append(tcVAlign)
    else:
        # Set vertical alignment
        if vertical_alignment == "top":
            cell.vertical_alignment = DOCX_CELL_ALIGN_VERTICAL.TOP
        elif vertical_alignment == "center":
            cell.vertical_alignment = DOCX_CELL_ALIGN_VERTICAL.CENTER
        elif vertical_alignment == "bottom":
            cell.vertical_alignment = DOCX_CELL_ALIGN_VERTICAL.BOTTOM
        else:
            raise ValueError(f"Invalid vertical alignment: {vertical_alignment}")


def _create_table_cell_subdocument(
    cell: DOCX_Cell, v_align="top", cell_type=None
) -> DOCX_Cell:
    """Create a subdocument for a table cell

    Args:
        cell: The table cell to use as a subdocument
        v_align: Vertical alignment ('top', 'center', 'bottom')
        cell_type: Type of cell ('about', 'contact', 'sidebar', 'main') for custom margins

    Returns:
        The prepared cell object
    """
    # Clear any default paragraph if present
    if len(cell.paragraphs) > 0:
        p = cell.paragraphs[0]
        if not p.text.strip():  # Only remove if empty
            p._p.getparent().remove(p._p)

    # Get margin values from config based on cell type
    margin_top = ConfigHelper.get_style_constant(
        f"{cell_type}_cell_margin_top",
        ConfigHelper.get_style_constant("cell_margin_top", 3),
    )
    margin_bottom = ConfigHelper.get_style_constant(
        f"{cell_type}_cell_margin_bottom",
        ConfigHelper.get_style_constant("cell_margin_bottom", 3),
    )
    margin_left = ConfigHelper.get_style_constant(
        f"{cell_type}_cell_margin_left",
        ConfigHelper.get_style_constant("cell_margin_left", 3),
    )
    margin_right = ConfigHelper.get_style_constant(
        f"{cell_type}_cell_margin_right",
        ConfigHelper.get_style_constant("cell_margin_right", 3),
    )

    _cell_margins(cell, margin_top, margin_bottom, margin_left, margin_right)

    _cell_vertical_alignment(cell, v_align)

    return cell


def _apply_table_cell_fill_and_border_styles(
    cell: DOCX_Cell, styles: Dict[str, str | int | bool]
) -> None:
    """Apply fill and border styles to a table cell

    Args:
        cell: The cell object
        styles: Dictionary of styles to apply
    """

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    fill_enabled = styles.get("fill_enabled", True)

    if fill_enabled:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), styles.get("fill_color", "FFFFFF"))
        shading.set(qn("w:val"), "clear")
        tcPr.append(shading)

    border_enabled = styles.get("border_enabled", False)

    if border_enabled:
        tcBorders = OxmlElement("w:tcBorders")

        # Get border settings from config
        border_width = styles.get("border_width", 1)
        border_color = styles.get("border_color", "FFFFFF")

        for border in ["top", "left", "bottom", "right"]:
            if border == "top" and not styles.get("border_top_enabled", border_enabled):
                continue
            if border == "left" and not styles.get(
                "border_left_enabled", border_enabled
            ):
                continue
            if border == "bottom" and not styles.get(
                "border_bottom_enabled", border_enabled
            ):
                continue
            if border == "right" and not styles.get(
                "border_right_enabled", border_enabled
            ):
                continue
            border_elem = OxmlElement(f"w:{border}")
            border_elem.set(qn("w:val"), "single")  # 'single' for solid line
            if border == "top":
                border_width = styles.get("border_top_width", border_width)
                border_color = styles.get("border_top_color", border_color)
            elif border == "bottom":
                border_width = styles.get("border_bottom_width", border_width)
                border_color = styles.get("border_bottom_color", border_color)
            elif border == "left":
                border_width = styles.get("border_left_width", border_width)
                border_color = styles.get("border_left_color", border_color)
            elif border == "right":
                border_width = styles.get("border_right_width", border_width)
                border_color = styles.get("border_right_color", border_color)
            else:
                raise ValueError(f"Unknown border type: {border}")
            border_elem.set(
                qn("w:sz"), str(border_width * 8)
            )  # Size in eighths of a point
            border_elem.set(qn("w:color"), border_color)  # White border
            tcBorders.append(border_elem)

        tcPr.append(tcBorders)


def _create_two_column_layout(document: DOCX_Document) -> tuple:
    """Create the two-column layout structure with contact ribbon AFTER about section

    Args:
        document: The Word document

    Returns:
        tuple: (about_content_cell, contact_info_cell, sidebar_cell, main_cell)
    """
    # Get page width excluding margins
    page_width = ConfigHelper.get_document_defaults().get("page_width", 8.5)
    margin_left = ConfigHelper.get_document_defaults().get("margin_left", 0.2)
    margin_right = ConfigHelper.get_document_defaults().get("margin_right", 0.2)
    margins = margin_left + margin_right
    content_width = page_width - margins

    sidebar_width_ratio = ConfigHelper.get_document_defaults().get(
        "sidebar_width_ratio", 0.33
    )
    main_width_ratio = ConfigHelper.get_document_defaults().get(
        "main_width_ratio", 0.67
    )

    # Calculate column widths
    sidebar_width = content_width * sidebar_width_ratio
    main_width = content_width * main_width_ratio

    # Determine sidebar position (left or right)
    sidebar_on_right = (
        ConfigHelper.get_document_defaults().get("sidebar_position", "left") == "right"
    )

    about_styles = ConfigHelper.get_style_constant("about", {})

    # Create about section table FIRST (1 row, 1 column) for professional summary
    about_table = _create_table(document, 1, 1)

    about_content_cell = _create_table_cell_subdocument(
        about_table.cell(0, 0), cell_type="about"
    )

    _apply_table_cell_fill_and_border_styles(about_table.cell(0, 0), about_styles)

    # THEN create contact ribbon table (1 row, 1 column) AFTER about section
    contact_table = _create_table(document, 1, 1)

    contact_info_cell = _create_table_cell_subdocument(
        contact_table.cell(0, 0), v_align="center", cell_type="contact"
    )

    # Create the main content table (1 row, 2 columns)
    content_table = _create_table(document, 1, 2)

    # Get and prepare the cells with correct vertical alignment
    if sidebar_on_right:
        # Main content on left, sidebar on right
        content_table.columns[0].width = Inches(main_width)
        content_table.columns[1].width = Inches(sidebar_width)
        main_cell = _create_table_cell_subdocument(
            content_table.cell(0, 0), cell_type="main"
        )
        sidebar_cell = _create_table_cell_subdocument(
            content_table.cell(0, 1), v_align="top", cell_type="sidebar"
        )

        # Apply background color to the right cell
        cell_sidebar = content_table.cell(0, 1)
        cell_main = content_table.cell(0, 0)
    else:
        # Sidebar on left (default), main content on right
        content_table.columns[0].width = Inches(sidebar_width)
        content_table.columns[1].width = Inches(main_width)
        sidebar_cell = _create_table_cell_subdocument(
            content_table.cell(0, 0), v_align="top", cell_type="sidebar"
        )
        main_cell = _create_table_cell_subdocument(
            content_table.cell(0, 1), cell_type="main"
        )

        # Apply background color to the left cell
        cell_sidebar = content_table.cell(0, 0)
        cell_main = content_table.cell(0, 1)

    main_content_styles = ConfigHelper.get_style_constant("main_content", {})
    _apply_table_cell_fill_and_border_styles(cell_main, main_content_styles)

    sidebar_styles = ConfigHelper.get_style_constant("sidebar", {})
    _apply_table_cell_fill_and_border_styles(cell_sidebar, sidebar_styles)

    return about_content_cell, contact_info_cell, sidebar_cell, main_cell


def _add_horizontal_line_to_cell(cell: DOCX_Cell, is_sidebar: bool = False):
    """Add a simple horizontal line to a table cell

    Args:
        cell: The cell to add the line to
        is_sidebar: Whether this is in the sidebar (affects line length)
    """
    # Get values from config with fallbacks
    line_char = ConfigHelper.get_style_constant("horizontal_line_char", "_")
    line_spacing = ConfigHelper.get_style_constant("horizontal_line_spacing", 1.6)
    line_color = ConfigHelper.get_style_constant("horizontal_line_color", "000000")
    line_font_size = ConfigHelper.get_style_constant("horizontal_line_font_size", 16)

    if is_sidebar:
        line_length = ConfigHelper.get_style_constant(
            "sidebar_horizontal_line_length", 30
        )
    else:
        line_length = ConfigHelper.get_style_constant("horizontal_line_length", 50)

    p = cell.add_paragraph()
    _paragraph_alignment(p, "center")
    p.add_run(line_char * line_length)

    for run in p.runs:
        # Set font size for the horizontal line
        # FIXME: For some reason, correct line spacing is only applied when font is set to a value greater than 14 (H1)
        _apply_font_properties(
            run.font,
            {
                "bold": True,
                "color": line_color,
                "font_size": line_font_size,
            },
        )

    # Get specific header-to-about spacing if it exists
    if line_spacing:
        header_to_about = ConfigHelper.get_style_constant(
            "header_to_about_spacing", None
        )
        if header_to_about is not None:
            # Use specific header-to-about spacing
            _apply_paragraph_format_properties(
                p.paragraph_format, {"line_spacing": header_to_about}
            )
        else:
            # Use general line spacing
            _apply_paragraph_format_properties(
                p.paragraph_format, {"line_spacing": line_spacing}
            )


##############################
# Primary Helpers
##############################
def _apply_font_properties(font_obj: DOCX_FONT, properties: Dict[str, str]) -> None:
    """Apply font properties to a font object

    Args:
        font_obj: The font object (from style.font or run.font)
        properties: Dictionary containing font properties
    """
    if "font_name" in properties:
        font_obj.name = properties["font_name"]
    if "font_size" in properties:
        font_obj.size = Pt(properties["font_size"])
    if "bold" in properties:
        font_obj.bold = properties["bold"]
    if "italic" in properties:
        font_obj.italic = properties["italic"]
    if "underline" in properties:
        font_obj.underline = properties["underline"]
    if "color" in properties:
        font_obj.color.rgb = RGBColor.from_string(properties["color"])


def _apply_paragraph_format_properties(
    paragraph_format: DOCX_ParagraphFormat, properties: Dict[str, str | int | float]
) -> None:
    """Apply paragraph format properties to a paragraph format object

    Args:
        paragraph_format: The paragraph format object
        properties: Dictionary containing paragraph format properties
    """
    if "line_spacing" in properties and properties["line_spacing"] is not None:
        paragraph_format.line_spacing = properties["line_spacing"]
    if "space_before" in properties and properties["space_before"] is not None:
        paragraph_format.space_before = Pt(properties["space_before"])
    if "space_after" in properties and properties["space_after"] is not None:
        paragraph_format.space_after = Pt(properties["space_after"])
    if "indent_left" in properties and properties["indent_left"] is not None:
        paragraph_format.left_indent = Inches(properties["indent_left"])
    if "indent_right" in properties and properties["indent_right"] is not None:
        paragraph_format.right_indent = Inches(properties["indent_right"])
    if "alignment" in properties and properties["alignment"] is not None:
        paragraph_format.alignment = properties["alignment"]


def _prepare_section(
    document: DOCX_Document,
    soup: BeautifulSoup,
    section_type: ResumeSection,
    space_before: int | None = None,
    space_after: int | None = None,
) -> BS4_Element | None:
    """Universal preliminary section preparation

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content
        section_type: ResumeSection enum value

    Returns:
        BeautifulSoup element or None: The section heading element if found, None otherwise
    """
    section_h2 = soup.find("h2", string=lambda text: section_type.matches(text))

    if not section_h2:
        print(f"ℹ️  Section '{section_type.docx_heading}' not found in document")
        return None

    if not space_before:
        space_before = (
            section_type.space_before_h2 if section_type.add_space_before_h2 else None
        )
    if not space_after:
        space_after = section_type.space_after_h2

    # Check if there's an HR before this section
    section_page_break = _has_hr_before_element(section_h2)

    # Add page break if HR is found
    if section_page_break:
        # Get the last paragraph in the document
        if document.paragraphs:
            last_para = document.paragraphs[-1]
            # Add page break to the LAST paragraph instead of creating a new one
            if last_para.runs:
                last_para.runs[-1].add_break(DOCX_BREAK_TYPE.PAGE)
            else:
                run = last_para.add_run()
                run.add_break(DOCX_BREAK_TYPE.PAGE)
        else:
            # No paragraphs exist, create one with page break
            p = document.add_paragraph()
            run = p.add_run()
            run.add_break(DOCX_BREAK_TYPE.PAGE)

    # FIXME: This should only be enabled for two_column layouts
    # elif section_type.add_space_before_h2:
    #     _add_space_paragraph(document, ConfigHelper.get_style_constant("font_size_pts"))

    # Add the section heading
    use_paragraph_style = HeadingsHelper.should_use_paragraph_style("h2")
    heading_level = HeadingsHelper.get_level_for_tag("h2")
    _add_heading_or_paragraph(
        document,
        section_type.docx_heading,
        heading_level,
        use_paragraph_style=use_paragraph_style,
        space_before=space_before,
        space_after=space_after,
    )

    return section_h2


def _process_horizontal_skills_list(
    document: DOCX_Document,
    text: str,
    is_top_skills: bool = False,
    custom_input_separator: str = None,
    custom_output_separator: str = None,
) -> DOCX_Paragraph:
    """Process skills text into a horizontal list with consistent formatting

    Args:
        document: The Word document object
        text: Text containing skills separated by a delimiter
        is_top_skills: Whether this is the top skills section (affects styling)
        custom_input_separator: Custom input separator override
        custom_output_separator: Custom output separator override

    Returns:
        DOCX_Paragraph: The created paragraph
    """
    # Get configuration based on skills type
    if is_top_skills:
        config_prefix = "top_skills"
        default_input_sep = " • "
        default_output_sep = " | "
        use_formatted_paragraph = True
    else:
        config_prefix = "key_skills"
        default_input_sep = " · "
        default_output_sep = ", "
        use_formatted_paragraph = False

    # Get separators and formatting options
    input_separator = custom_input_separator or ConfigHelper.get_style_constant(
        f"{config_prefix}_separator_markdown", default_input_sep
    )
    output_separator = custom_output_separator or ConfigHelper.get_style_constant(
        f"{config_prefix}_separator", default_output_sep
    )
    apply_bold = ConfigHelper.get_style_constant(f"{config_prefix}_bold", False)

    # Get font size and line spacing (key skills specific)
    font_size = ConfigHelper.get_style_constant(f"{config_prefix}_font_size", None)
    line_spacing = ConfigHelper.get_style_constant(
        f"{config_prefix}_line_spacing", None
    )

    # Parse skills from text
    skills = [s.strip() for s in text.split(input_separator.strip())]
    skills = [s for s in skills if s]

    # Format and add to document
    paragraph = _format_skills_list(
        document, skills, output_separator, apply_bold, use_formatted_paragraph
    )

    # Apply additional formatting if specified and not using formatted paragraph
    if not use_formatted_paragraph:
        # Apply font size to all runs in the paragraph
        if font_size:
            for run in paragraph.runs:
                _apply_font_properties(run.font, {"font_size": font_size})

        # Apply line spacing to the paragraph
        if line_spacing:
            _apply_paragraph_format_properties(
                paragraph.paragraph_format, {"line_spacing": line_spacing}
            )

    return paragraph


def _process_tagline_and_specialty(
    container: DOCX_Document, first_p: BS4_Element, alignment_str: str = "center"
) -> None:
    """Process tagline and specialty paragraphs

    Args:
        container: Document or cell to add content to
        first_p: First paragraph element (tagline)
        alignment_str: Alignment as string (left, right, center)
    """
    # Check if ANY paragraph headings are specified
    use_paragraph_style = HeadingsHelper.any_heading_uses_paragraph_style()

    # Check if the paragraph contains emphasis (italics)
    em_tag = first_p.find("em")
    if em_tag:
        # Process paragraph with emphasis
        if use_paragraph_style:
            tagline_para = container.add_paragraph()
            tagline_run = tagline_para.add_run(em_tag.text)
            _apply_font_properties(
                tagline_run.font,
                {
                    "italic": True,
                    "font_size": HeadingsHelper.get_font_size_for_tag("h4"),
                },
            )
        else:
            tagline_para = container.add_paragraph(em_tag.text, style="Subtitle")
            for run in tagline_para.runs:
                run.italic = True
                _apply_font_properties(
                    run.font,
                    {
                        "italic": True,
                    },
                )

        _paragraph_alignment(tagline_para, alignment_str)

        # Add the rest of the paragraph if any
        rest_of_p = first_p.text.replace(em_tag.text, "").strip()
        if rest_of_p:
            rest_para = container.add_paragraph()
            _paragraph_alignment(rest_para, alignment_str)
            rest_para.add_run(rest_of_p)
    else:
        # Process regular paragraph
        if use_paragraph_style:
            tagline_para = container.add_paragraph()
            _paragraph_alignment(tagline_para, alignment_str)
            tagline_para.add_run(first_p.text)
        else:
            tagline_para = container.add_paragraph(first_p.text, style="Subtitle")
            _paragraph_alignment(tagline_para, alignment_str)

    # Process additional specialty paragraphs
    current_p = first_p.find_next_sibling()
    while current_p and current_p.name == "p":
        specialty_para = container.add_paragraph()
        _paragraph_alignment(specialty_para, alignment_str)
        specialty_para.add_run(current_p.text)

        # Apply spacing
        header_specialty_space_after = ConfigHelper.get_style_constant(
            "header_specialty_space_after", None
        )
        header_specialty_space_before = ConfigHelper.get_style_constant(
            "header_specialty_space_before", None
        )
        _apply_paragraph_format_properties(
            specialty_para.paragraph_format,
            {
                "space_after": header_specialty_space_after,
                "space_before": header_specialty_space_before,
            },
        )

        current_p = current_p.find_next_sibling()


def _process_simple_section(
    document: DOCX_Document,
    section_h2: BS4_Element,
    add_space: bool = False,
) -> None:
    """Process sections with simple paragraph-based content like Education and Contact.
    These sections typically have paragraphs with some bold (strong) elements.

    Args:
        document: The Word document object
        section_h2: The BeautifulSoup h2 element for the section
        add_space (bool, optional): Whether to add a space paragraph after the section. Defaults to False.

    Returns:
        None
    """
    if not section_h2:
        return

    current_element = section_h2.find_next_sibling()

    while current_element and current_element.name != "h2":
        if current_element.name == "p":
            para = document.add_paragraph()
            for child in current_element.children:
                if getattr(child, "name", None) == "strong":
                    para.add_run(f"{child.text}: ").bold = True
                else:
                    if child.string and child.string.strip():
                        _process_text_for_hyperlinks(para, child.string.strip())
        elif current_element.name == "ul":
            # Handle bullet lists if they appear
            _add_bullet_list(document, current_element)

        current_element = current_element.find_next_sibling()

    # Add an extra space after the section if requested
    if add_space:
        _add_space_paragraph(document, ConfigHelper.get_style_constant("font_size_pts"))


def _process_project_section(
    document: DOCX_Document,
    project_element: BS4_Element,
    processed_elements: set[BS4_Element],
) -> set[BS4_Element]:
    """Process a project/client section and its related elements

    Args:
        document: The Word document object
        project_element: BeautifulSoup element for the project heading
        processed_elements: Set of elements already processed

    Returns:
        set: Updated set of processed elements
    """
    bullet_indent_inches = ConfigHelper.get_style_constant("bullet_indent_inches")

    # Get the next element to see if it contains the project details
    next_element = project_element.find_next_sibling()
    project_info = ""

    # If next element is a paragraph, it might contain the project name/duration
    if next_element and next_element.name == "p":
        project_info = next_element.text.strip()
        processed_elements.add(next_element)
        next_element = next_element.find_next_sibling()

    # Get the proper subsection
    subsection = JobSubsection.find_by_tag_and_text(
        project_element.name, project_element.text
    )
    if not subsection:
        subsection = JobSubsection.PROJECT_CLIENT  # Fallback

    heading_level = HeadingsHelper.get_level_for_tag(project_element.name)
    use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
        project_element.name
    )

    # Prepare the project text
    project_text = subsection.full_heading
    if project_info:
        project_text += " " + project_info

    # Add the heading or formatted paragraph
    _add_heading_or_paragraph(
        document,
        project_text,
        heading_level,
        use_paragraph_style=use_paragraph_style,
        bold=subsection.bold,
        italic=subsection.italic,
    )

    # Process next elements under this project until another section
    while (
        next_element
        and next_element.name not in ["h2", "h3", "h4"]
        and not (next_element.name == "h5")
    ):
        if next_element.name == "p" and not next_element.find("h6"):
            # Process regular paragraph text
            para = document.add_paragraph()
            _process_text_for_hyperlinks(para, next_element.text.strip())
            _left_indent_paragraph(para, bullet_indent_inches / 2)
            processed_elements.add(next_element)
            next_element = next_element.find_next_sibling()
            continue

        # Find subsection for h6 elements
        h6_subsection = None
        if next_element.name == "h6":
            h6_subsection = JobSubsection.find_by_tag_and_text(
                next_element.name, next_element.text
            )

        # Responsibilities Overview
        if h6_subsection == JobSubsection.RESPONSIBILITIES:
            heading_level = HeadingsHelper.get_level_for_tag(next_element.name)
            use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
                next_element.name
            )

            # Add the heading or paragraph
            if use_paragraph_style:
                resp_header = document.add_paragraph()
                _left_indent_paragraph(resp_header)  # Keep indentation
                resp_run = resp_header.add_run(h6_subsection.full_heading)
                resp_run.bold = h6_subsection.bold
                resp_run.italic = h6_subsection.italic
            else:
                resp_heading = document.add_heading(
                    h6_subsection.full_heading, level=heading_level
                )
                _left_indent_paragraph(resp_heading)  # Keep indentation

            # Get the paragraph with responsibilities
            resp_element = next_element.find_next_sibling()
            if resp_element and resp_element.name == "p":
                resp_para = document.add_paragraph()
                _process_text_for_hyperlinks(resp_para, resp_element.text)
                _left_indent_paragraph(resp_para)
                processed_elements.add(resp_element)

            processed_elements.add(next_element)

        # Additional Details
        elif h6_subsection == JobSubsection.ADDITIONAL_DETAILS:
            heading_level = HeadingsHelper.get_level_for_tag(next_element.name)
            use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
                next_element.name
            )

            # Add the heading or paragraph
            if use_paragraph_style:
                details_header = document.add_paragraph()
                _left_indent_paragraph(details_header)  # Keep indentation
                details_run = details_header.add_run(h6_subsection.full_heading)
                details_run.bold = h6_subsection.bold
                details_run.italic = h6_subsection.italic
            else:
                details_heading = document.add_heading(
                    h6_subsection.full_heading, level=heading_level
                )
                _left_indent_paragraph(details_heading)  # Keep indentation

            processed_elements.add(next_element)

        # Bullet points
        elif next_element.name == "ul":
            _left_indent_paragraph(_add_bullet_list(document, next_element))
            processed_elements.add(next_element)

        next_element = next_element.find_next_sibling()

    return processed_elements


def _process_job_entry(
    document: DOCX_Document,
    job_element: BS4_Element,
    processed_elements: set[BS4_Element],
    space_before: int | None = None,
    space_after: int | None = None,
    is_first_job: bool = True,
) -> set[BS4_Element]:
    """Process a job entry (h3) and its related elements"""
    job_title = job_element.text.strip()

    # Check for HR before h3 and add page break if found
    if _has_hr_before_element(job_element):
        # Get config option for blank line - default to True if not specified
        hr_blank_line_enabled = ConfigHelper.get_style_constant(
            "hr_before_job_blank_line", True
        )

        if hr_blank_line_enabled:
            p = document.add_paragraph()
            run = p.add_run()
            run.add_break(DOCX_BREAK_TYPE.PAGE)
    # Add horizontal line before job entries (except the first one)
    elif not is_first_job and hasattr(
        document, "tables"
    ):  # Check if it's a cell/subdocument
        # Get values from config with fallbacks
        line_char = ConfigHelper.get_style_constant("horizontal_line_char", "_")
        line_length = ConfigHelper.get_style_constant(
            "job_entry_horizontal_line_length", 30
        )

        # Add the horizontal line
        p = document.add_paragraph()
        _paragraph_alignment(p, "center")
        p.add_run(line_char * line_length)

        for run in p.runs:
            # Set font size for the horizontal line
            _apply_font_properties(
                run.font,
                {
                    "bold": True,
                    "font_size": Pt(16),
                },
            )

        _apply_paragraph_format_properties(
            p.paragraph_format,
            {
                "line_spacing": ConfigHelper.get_style_constant(
                    "job_entry_horizontal_line_spacing", 1.0
                )
            },
        )

    # Add explicit space before h3 in two-column mode by adding an empty paragraph
    elif space_before and not is_first_job:
        # Add a spacer paragraph with specific height before company name
        spacer = document.add_paragraph()
        spacer.add_run(" ")  # Need at least one character

        spacer.paragraph_format.space_after = Pt(space_before)

    # Add the company name as h3
    use_paragraph_style = HeadingsHelper.should_use_paragraph_style("h3")
    heading_level = HeadingsHelper.get_level_for_tag("h3")
    company_heading = _add_heading_or_paragraph(
        document,
        job_title,
        heading_level,
        use_paragraph_style=use_paragraph_style,
        # Don't pass space_before here since we handled it manually above
        space_after=space_after,
    )

    # Mark the h3 as processed
    processed_elements.add(job_element)

    # Check for the paragraph immediately after h3
    next_element = job_element.find_next_sibling()

    # Process paragraph with duration text if it exists
    if next_element and next_element.name == "p":
        duration_para = document.add_paragraph()

        # Check for both bold and italic formatting
        has_strong = next_element.find("strong")
        has_em = next_element.find("em")

        duration_text = next_element.text.replace("*", "").replace("_", "").strip()
        duration_run = duration_para.add_run(duration_text)

        if has_strong:
            duration_run.bold = True
        if has_em:
            duration_run.italic = True

        processed_elements.add(next_element)

    return processed_elements


def _process_subsection(
    document: DOCX_Document,
    current_element: BS4_Element,
    subsection: JobSubsection,
    heading_level: int,
    processed_elements: set[BS4_Element],
) -> set[BS4_Element]:
    """Generic function to process any subsection (Summary, Internal, Responsibilities, etc.)

    Args:
        document: The Word document object
        current_element: The subsection heading element
        subsection: The JobSubsection enum value
        heading_level: The heading level from HeadingsHelper
        processed_elements: Set of elements already processed

    Returns:
        Updated set of processed elements
    """
    if current_element not in processed_elements:

        processed_elements.add(current_element)

        # Add the subsection heading
        use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
            current_element.name
        )
        _add_heading_or_paragraph(
            document,
            subsection.full_heading,
            heading_level,
            use_paragraph_style=use_paragraph_style,
            bold=subsection.bold,
            italic=subsection.italic,
        )

        # Process elements under this subsection until we hit another heading
        next_element = current_element.find_next_sibling()
        stop_tags = ["h2", "h3", "h4", "h5", "h6"]

        while next_element and next_element.name not in stop_tags:
            if next_element.name == "p":
                para = document.add_paragraph()
                _process_text_for_hyperlinks(para, next_element.text.strip())
                processed_elements.add(next_element)
            elif next_element.name == "ul":
                _add_bullet_list(document, next_element)
                processed_elements.add(next_element)

            next_element = next_element.find_next_sibling()

    return processed_elements


def _process_position(
    document: DOCX_Document,
    element: BS4_Element,
    processed_elements: set[BS4_Element],
    space_before: int | None = None,
    space_after: int | None = None,
) -> set[BS4_Element]:
    """Process a position entry (h4) and its related elements

    Args:
        document: The Word document object
        job: BeautifulSoup element
        processed_elements: Set of elements already processed
        space_before: Whether to add space before h4 headings
        space_after: Space after h4 heading, if any

    Returns:
        Updated set of processed elements
    """
    position_title = element.text.strip()

    # Add the position heading
    heading_level = HeadingsHelper.get_level_for_tag(element.name)
    use_paragraph_style = HeadingsHelper.should_use_paragraph_style(element.name)
    position_para = _add_heading_or_paragraph(
        document,
        position_title,
        heading_level,
        use_paragraph_style=use_paragraph_style,
        space_before=space_before,
        space_after=space_after,
    )

    _apply_paragraph_format_properties(
        position_para.paragraph_format,
        {
            "line_spacing": ConfigHelper.get_style_constant(
                "position_title_line_spacing", None
            )
        },
    )

    processed_elements.add(element)
    next_element = element.find_next_sibling()

    # Process date and location elements (either h6 or p format)
    if next_element:
        # Handle paragraph date/location format
        # Handle paragraph with combined date and location on single line
        if next_element.name == "p":
            if next_element.find("strong") or next_element.find("em"):

                # Check if next paragraph is location
                next_next_element = next_element.find_next_sibling()
                if (
                    next_next_element
                    and next_next_element.name == "p"
                    and (
                        next_next_element.find("strong") or next_next_element.find("em")
                    )
                ):
                    # Create a single paragraph with both date and location
                    date_text = (
                        next_element.text.replace("*", "").replace("_", "").strip()
                    )
                    location_text = (
                        next_next_element.text.replace("*", "").replace("_", "").strip()
                    )

                    reduced_spacing = float(
                        ConfigHelper.get_style_constant(
                            "date_location_line_spacing", 1.2
                        )
                    )
                    combo_para = document.add_paragraph()
                    _apply_paragraph_format_properties(
                        combo_para.paragraph_format,
                        {
                            "line_spacing": reduced_spacing,
                        },
                    )

                    # Get font size from config
                    date_loc_font_size = ConfigHelper.get_style_constant(
                        "date_location_font_size", None
                    )

                    # Get separator from config (default to forward slash if not specified)
                    date_location_separator = ConfigHelper.get_style_constant(
                        "date_location_separator", " / "
                    )

                    # Add date with appropriate formatting
                    date_run = combo_para.add_run(date_text + date_location_separator)
                    _apply_font_properties(
                        date_run.font,
                        {
                            "bold": next_element.find("strong") is not None,
                            "italic": next_element.find("em") is not None,
                            "font_size": date_loc_font_size,
                        },
                    )

                    # Add location with appropriate formatting
                    location_run = combo_para.add_run(location_text)
                    _apply_font_properties(
                        location_run.font,
                        {
                            "bold": next_element.find("strong") is not None,
                            "italic": next_element.find("em") is not None,
                            "font_size": date_loc_font_size,
                        },
                    )

                    processed_elements.add(next_element)
                    processed_elements.add(next_next_element)
                    next_element = next_next_element.find_next_sibling()

    return processed_elements


def _add_heading_or_paragraph(
    document: DOCX_Document,
    text: str,
    heading_level: int,
    use_paragraph_style: bool = False,
    bold: bool = True,
    italic: bool = False,
    font_size: int | None = None,
    space_before: int | None = None,
    space_after: int | None = None,
) -> DOCX_Paragraph:
    """Add either a heading or a formatted paragraph based on preference

    Args:
        document: The Word document object
        text (str): Text content for the heading/paragraph
        heading_level (int): The heading level (0-5) to use if not using paragraph style
        use_paragraph_style (bool): Whether to use paragraph style instead of heading style
        bold (bool): Whether to make the paragraph text bold (if using paragraph style)
        italic (bool): Whether to make the paragraph text italic (if using paragraph style)
        font_size (int, optional): Font size in points for paragraph style. Defaults to None.

    Returns:
        The created heading or paragraph object
    """
    has_add_heading = hasattr(document, "add_heading")
    format_styles = {}

    if use_paragraph_style or not has_add_heading:
        # For cells or when paragraph style is requested, use add_paragraph
        para = document.add_paragraph()
        run = para.add_run(text)
        format_styles["bold"] = bold
        format_styles["italic"] = italic

        # Apply appropriate font size
        if font_size is None:
            size_pt = HeadingsHelper.get_font_size_for_level(heading_level)
            format_styles["font_size"] = size_pt
        else:
            format_styles["font_size"] = font_size

        _apply_font_properties(run.font, format_styles)
    else:
        # Only use add_heading when available (for Document objects)
        para = document.add_heading(text, level=heading_level)

    _add_space_before_or_after(
        para,
        space_before,
        space_after,
    )

    return para


def _add_bullet_list(
    document: DOCX_Document, ul_element: BS4_Element, indentation: float = None
) -> DOCX_Paragraph:
    """Add bullet points from an unordered list element

    Args:
        document: The Word document object
        ul_element: BeautifulSoup element containing the unordered list
        indentation (float, optional): Left indentation in inches

    Returns:
        paragraph: The last bullet paragraph added
    """
    use_paragraph_style = ConfigHelper.get_style_constant("paragraph_lists", False)

    if use_paragraph_style:
        # Get the bullet character from config
        bullet_char = ConfigHelper.get_paragraph_list_option("ul", "bullet_character")

        para = document.add_paragraph()

        items = ul_element.find_all("li")
        for i, li in enumerate(items):
            if i > 0:
                para.add_run("\n")

            # Add the bullet character first
            para.add_run(f"{bullet_char} ")

            # Process formatting using the helper function
            # Always ensure sentence ending for each line in paragraph lists
            _process_list_item_formatting(para, li, ensure_ending=True)

        if indentation:
            _left_indent_paragraph(para, indentation)

        return para
    else:
        # Standard bullet list processing - unchanged
        bullet_para = None
        for li in ul_element.find_all("li"):
            bullet_para = document.add_paragraph(style="List Bullet")

            # Process formatting using the helper function
            _process_list_item_formatting(bullet_para, li)

            # Ensure sentence ending for each bullet
            last_run = bullet_para.runs[-1] if bullet_para.runs else None
            if last_run and not last_run.text.rstrip()[-1:] in [
                ".",
                "!",
                "?",
                ":",
                ";",
            ]:
                last_run.text = last_run.text.rstrip() + "."

            if indentation:
                _left_indent_paragraph(bullet_para, indentation)

        return bullet_para


def _process_projects_or_certifications(
    document: DOCX_Document,
    section_h2: BeautifulSoup,
    space_before_h3: int | None = None,
    space_after_h3: int | None = None,
    space_before_h4: int | None = None,
    space_after_h4: int | None = None,
) -> None:
    """Process the certifications section with its specific structure

    Args:
        document: The Word document object
        section_h2: BeautifulSoup h2 element for the section

    Returns:
        None
    """
    if not section_h2:
        return

    current_element = section_h2.find_next_sibling()
    first_h3_after_h2 = True  # Track the first h3 after h2

    while current_element and current_element.name != "h2":
        # Process certification name (h3)
        if current_element.name == "h3":
            cert_name = current_element.text.strip()
            use_paragraph_style = HeadingsHelper.should_use_paragraph_style("h3")
            heading_level = HeadingsHelper.get_level_for_tag("h3")

            para = _add_heading_or_paragraph(
                document,
                cert_name,
                heading_level,
                use_paragraph_style=use_paragraph_style,
                # space_before=space_before,
                # space_after=space_after,
            )

            if space_before_h3 and not first_h3_after_h2:
                _add_space_before_or_after(
                    para,
                    space_before=space_before_h3,
                )

            else:
                first_h3_after_h2 = False  # After first h3 is processed

            _add_space_before_or_after(
                para,
                space_after=space_after_h3,
            )

            # Look for next elements - either blockquote or organization info directly
            next_element = current_element.find_next_sibling()

            # Process any paragraph text that comes after h3 but before blockquote
            while next_element and next_element.name == "p":
                # Add the paragraph text
                para = document.add_paragraph()
                _process_text_for_hyperlinks(para, next_element.get_text().strip())

                # Move to the next element
                next_element = next_element.find_next_sibling()

            # Handle blockquote (optional)
            if next_element and next_element.name == "blockquote":
                # Process the blockquote contents
                _process_project_or_certification_blockquote(
                    document, next_element, space_before_h4, space_after_h4
                )

            # If no blockquote, look for organization info directly
            elif next_element:
                # Try to find organization info (could be bold text or heading)
                if next_element.name in ["h4", "h5", "h6", "p"] and next_element.find(
                    "strong"
                ):
                    # Extract organization text
                    org_text = next_element.find("strong").text.strip()
                    org_para = document.add_paragraph()
                    org_run = org_para.add_run(org_text)
                    org_run.bold = True

                    # Look for date information in the next element
                    date_element = next_element.find_next_sibling()
                    if date_element and date_element.find("em"):
                        em_tag = date_element.find("em")
                        date_text = em_tag.text.strip()
                        date_para = document.add_paragraph()

                        # Check if the date is hyperlinked
                        parent_a = em_tag.find_parent("a")
                        if parent_a and parent_a.get("href"):
                            _add_hyperlink(date_para, date_text, parent_a["href"])
                        else:
                            date_run = date_para.add_run(date_text)
                            date_run.italic = True

            # Add spacing after each certification
            # _add_space_paragraph(document)

        current_element = current_element.find_next_sibling()


def _process_project_or_certification_blockquote(
    document: DOCX_Document,
    blockquote: BS4_Element,
    space_before_h4: int | None = None,
    space_after_h4: int | None = None,
) -> None:
    """Process the contents of a certification blockquote

    Args:
        document: The Word document object
        blockquote: BeautifulSoup blockquote element containing certification details

    Returns:
        None
    """
    # Find and process organization info
    org_element, _ = _find_organization_element(blockquote)

    if org_element:
        if org_element.name in ["h4", "h5", "h6"]:
            _create_heading_with_formatting_preservation(
                document,
                org_element,
                space_before=space_before_h4,
                space_after=space_after_h4,
            )
        else:  # paragraph with strong tag
            strong_tag = org_element.find("strong")
            if strong_tag:
                org_para = document.add_paragraph()
                org_run = org_para.add_run(strong_tag.text.strip())
                org_run.bold = True

    # Process all other content
    for item in blockquote.contents:
        if item == org_element or (isinstance(item, str) and not item.strip()):
            continue

        if (
            hasattr(item, "name")
            and item.name
            and item.name.startswith("h")
            and item != org_element
        ):
            _create_heading_with_formatting_preservation(document, item)

        elif hasattr(item, "name") and item.name == "p" and item != org_element:
            if not _process_date_paragraph(document, item):
                # Regular paragraph
                para = document.add_paragraph()

                # Check if this paragraph contains a Key Skills heading
                # Look for h5 elements with "key skills" text (case insensitive)
                has_key_skills_heading = False
                for child in item.children:
                    if (
                        getattr(child, "name", None) == "h5"
                        and child.text.strip().lower() == "key skills"
                    ):
                        has_key_skills_heading = True
                        break

                _process_element_children_with_formatting(
                    para, item, add_colon_to_strong=has_key_skills_heading
                )

        elif hasattr(item, "name") and item.name == "ul":
            _add_bullet_list(document, item)

        elif isinstance(item, str) and item.strip():
            para = document.add_paragraph()
            _process_text_for_hyperlinks(para, item.strip())


def _process_element_children_with_formatting(
    paragraph: DOCX_Paragraph, element: BS4_Element, add_colon_to_strong: bool = False
) -> None:
    """Process child elements of an HTML element and add them to a paragraph with proper formatting

    Args:
        paragraph: The Word paragraph to add content to
        element: BeautifulSoup element whose children to process
        add_colon_to_strong: Whether to add a colon after strong elements
    """
    for child in element.children:
        if getattr(child, "name", None) == "strong":
            text = child.text + (":" if add_colon_to_strong else "")
            run = paragraph.add_run(text)
            run.bold = True
        elif getattr(child, "name", None) == "em":
            run = paragraph.add_run(child.text)
            run.italic = True
        elif getattr(child, "name", None) == "a" and child.get("href"):
            _add_hyperlink(paragraph, child.text, child.get("href"))
        elif child.string:
            _process_text_for_hyperlinks(paragraph, child.string)


def _create_heading_with_formatting_preservation(
    document: DOCX_Document,
    element: BS4_Element,
    heading_level: int | None = None,
    use_paragraph_style: bool = None,
    space_before: int | None = None,
    space_after: int | None = None,
) -> None:
    """Create a heading or paragraph from an element while preserving child formatting

    Args:
        document: The Word document object
        element: BeautifulSoup element to process
        heading_level: Heading level to use (if None, will be determined from element)
        use_paragraph_style: Whether to use paragraph style (if None, will be determined)
    """
    if heading_level is None:
        heading_level = HeadingsHelper.get_level_for_tag(element.name)
    if use_paragraph_style is None:
        use_paragraph_style = HeadingsHelper.should_use_paragraph_style(element.name)

    # Check if document supports add_heading
    has_add_heading = hasattr(document, "add_heading")

    if use_paragraph_style or not has_add_heading:
        # For cells or when paragraph style is requested, use add_paragraph
        para = document.add_paragraph()
        # For cells that were meant to be headings, apply heading formatting manually
        if not use_paragraph_style and not has_add_heading:
            # Get font size for this heading level
            HeadingsHelper.get_font_size_for_level(heading_level)
            # We'll apply formatting to child elements below
    else:
        # Only use add_heading when available (for Document objects)
        para = document.add_heading(level=heading_level)

    # Process child elements with proper formatting
    _process_element_children_with_formatting(para, element)

    # For cell "headings", make sure all runs are bold
    if not use_paragraph_style and not has_add_heading:
        for run in para.runs:
            _apply_font_properties(
                run.font,
                {
                    "bold": True,
                    "font_size": HeadingsHelper.get_font_size_for_level(heading_level),
                },
            )

    # Apply spacing
    _add_space_before_or_after(
        para,
        space_before,
        space_after,
    )


def _process_date_paragraph(
    document: DOCX_Document, paragraph_element: BS4_Element
) -> bool:
    """Process a paragraph that contains date information (em tags)

    Args:
        document: The Word document object
        paragraph_element: BeautifulSoup paragraph element

    Returns:
        bool: True if date was processed, False otherwise
    """
    em_tag = paragraph_element.find("em")
    if not em_tag:
        return False

    date_text = em_tag.text.strip()
    date_para = document.add_paragraph()

    # Check if the date is hyperlinked
    parent_a = em_tag.find_parent("a")
    if parent_a and parent_a.get("href"):
        _add_hyperlink(date_para, date_text, parent_a["href"])
    else:
        date_run = date_para.add_run(date_text)
        date_run.italic = True

    return True


def _add_configured_page_numbers(document: DOCX_Document) -> None:
    """Add page numbers based on configuration settings

    Args:
        document: The Word document object
    """
    page_numbers_enabled = ConfigHelper.get_document_defaults().get(
        "page_numbers_enabled", False
    )
    if not page_numbers_enabled:
        return

    page_numbers_alignment = ConfigHelper.get_document_defaults().get(
        "page_numbers_alignment", "center"
    )

    section = document.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]

    _paragraph_alignment(footer_para, page_numbers_alignment)

    # Clear existing content
    footer_para.clear()

    # Get font settings from configuration
    font_size = ConfigHelper.get_style_constant("page_number_font_size", 10)
    font_name = ConfigHelper.get_style_constant("page_number_font_name", "Calibri")

    page_numbers_format = ConfigHelper.get_document_defaults().get(
        "page_numbers_format", "simple"
    )

    page_numbers_text = ConfigHelper.get_document_defaults().get(
        "page_numbers_text", "{page} of {total}"
    )

    if page_numbers_format == "simple":
        # Simple page numbers
        run = footer_para.add_run()
        _apply_font_properties(
            run.font,
            {
                "font_size": font_size,
                "font_name": font_name,
            },
        )

        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    elif page_numbers_format == "with_total":
        # Page X of Y format
        format_text = page_numbers_text
        parts = format_text.split("{page}")

        if len(parts) == 2:
            before_page, after_page_part = parts
            after_page_parts = after_page_part.split("{total}")

            if len(after_page_parts) == 2:
                between_page_total, after_total = after_page_parts

                # Add text before page number
                if before_page:
                    run = footer_para.add_run(before_page)
                    _apply_font_properties(
                        run.font,
                        {
                            "font_size": font_size,
                            "font_name": font_name,
                        },
                    )

                # Add current page number field
                run1 = footer_para.add_run()
                _apply_font_properties(
                    run1.font,
                    {
                        "font_size": font_size,
                        "font_name": font_name,
                    },
                )

                fldChar1 = OxmlElement("w:fldChar")
                fldChar1.set(qn("w:fldCharType"), "begin")
                instrText1 = OxmlElement("w:instrText")
                instrText1.text = "PAGE"
                fldChar2 = OxmlElement("w:fldChar")
                fldChar2.set(qn("w:fldCharType"), "end")
                run1._r.append(fldChar1)
                run1._r.append(instrText1)
                run1._r.append(fldChar2)

                # Add text between page and total
                if between_page_total:
                    run = footer_para.add_run(between_page_total)
                    _apply_font_properties(
                        run.font,
                        {
                            "font_size": font_size,
                            "font_name": font_name,
                        },
                    )

                # Add total pages field
                run2 = footer_para.add_run()
                _apply_font_properties(
                    run2.font,
                    {
                        "font_size": font_size,
                        "font_name": font_name,
                    },
                )

                fldChar3 = OxmlElement("w:fldChar")
                fldChar3.set(qn("w:fldCharType"), "begin")
                instrText2 = OxmlElement("w:instrText")
                instrText2.text = "NUMPAGES"
                fldChar4 = OxmlElement("w:fldChar")
                fldChar4.set(qn("w:fldCharType"), "end")
                run2._r.append(fldChar3)
                run2._r.append(instrText2)
                run2._r.append(fldChar4)

                # Add text after total
                if after_total:
                    run = footer_para.add_run(after_total)
                    _apply_font_properties(
                        run.font,
                        {
                            "font_size": font_size,
                            "font_name": font_name,
                        },
                    )


##############################
# Interactive Mode Helper
##############################
def _run_interactive_mode() -> tuple[Path, Path, dict[str, bool], bool, ConfigLoader]:
    """Run in interactive mode, prompting the user for inputs

    Returns:
        tuple: (input_file, output_file, paragraph_style_headings, create_pdf, config_loader)
    """
    print("\n🎯 Welcome to Resume Markdown to ATS Converter (Interactive Mode) 🎯\n")

    # Prompt for config file first
    default_config = DEFAULT_CONFIG_FILE
    config_prompt = (
        f"⚙️ Enter path to configuration file (default: '{default_config}'): "
    )
    config_file = input(config_prompt).strip()
    if not config_file:
        config_file = default_config
        print(f"✅ Using default configuration: {config_file}")

    # Create ConfigLoader with the specified config file
    config_loader = ConfigLoader(config_file)

    # Prompt for input file
    while True:
        input_file = input("📄 Enter the path to your Markdown resume file: ").strip()
        if not input_file:
            print("❌ Input file path cannot be empty. Please try again.")
            continue

        if not os.path.exists(input_file):
            print(f"❌ File '{input_file}' does not exist. Please enter a valid path.")
            continue

        break

    input_file = Path(input_file)

    # Get output file
    output_file = OutputFilePath(input_file=input_file, interactive=True).output_path()

    # Prompt for paragraph style headings
    print(
        "\n🔠 Choose heading levels to render as paragraphs instead of Word headings:"
    )
    print("   (Enter numbers separated by space, e.g., '3 4 5 6' for h3, h4, h5, h6)")
    print("   1. h3 - Job titles")
    print("   2. h4 - Company names")
    print("   3. h5 - Subsections (Key Skills, Summary, etc.)")
    print("   4. h6 - Sub-subsections (Responsibilities, Additional Details)")
    print("   0. None (use Word heading styles for all)")

    heading_choices = input(
        "👉 Your choices (e.g., '3 4 5 6' or '0' for none): "
    ).strip()

    paragraph_style_headings = {}
    if heading_choices != "0":
        chosen_numbers = [int(n) for n in heading_choices.split() if n.isdigit()]
        heading_map = {1: "h3", 2: "h4", 3: "h5", 4: "h6"}

        for num in chosen_numbers:
            if 1 <= num <= 4:
                heading_tag = heading_map[num]
                paragraph_style_headings[heading_tag] = True
                print(f"✅ Selected {heading_tag} for paragraph styling")
    else:
        print("✅ Using Word headings for all levels (no paragraph styling)")

    print("\n⚙️ Processing your resume...\n")

    create_pdf = (
        input("📄 Also create a PDF version? (y/n, default: n): ").strip().lower()
        == "y"
    )
    if create_pdf:
        print("✅ Will generate PDF output")

    # Return the ConfigLoader object directly instead of just the config_file path
    return input_file, output_file, paragraph_style_headings, create_pdf, config_loader


##############################
# PDF Helpers
##############################
def _convert_with_docx2pdf(docx_file: str, pdf_file: str) -> bool:
    """Convert using docx2pdf library"""
    try:
        from docx2pdf import convert

        convert(docx_file, pdf_file)
        return os.path.exists(pdf_file)
    except ImportError:
        return False


def _convert_with_libreoffice(docx_file: str, pdf_file: str) -> bool:
    """Convert using LibreOffice command line"""
    import platform
    import subprocess

    # Find the LibreOffice executable based on platform
    if platform.system() == "Windows":
        for path in PdfConverterPaths.LIBREOFFICE_WINDOWS.value:
            if os.path.exists(path):
                lo_exec = path
                break
        else:
            return False
    elif platform.system() == "Darwin":  # macOS
        lo_exec = PdfConverterPaths.LIBREOFFICE_MACOS.value
        if not os.path.exists(lo_exec):
            return False
    else:  # Linux/Unix
        lo_exec = PdfConverterPaths.LIBREOFFICE_LINUX.value  # Assume it's in PATH

    try:
        subprocess.run(
            [
                lo_exec,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                os.path.dirname(pdf_file) or ".",
                docx_file,
            ],
            check=True,
            stdout=subprocess.DEVNULL,
        )

        # LibreOffice saves to the original filename with .pdf extension
        temp_pdf = (
            os.path.splitext(os.path.basename(docx_file))[0] + f".{PDF_EXTENSION}"
        )
        temp_pdf_path = os.path.join(os.path.dirname(pdf_file) or ".", temp_pdf)

        # If the output path is different from LibreOffice's default, move the file
        if os.path.abspath(temp_pdf_path) != os.path.abspath(pdf_file):
            os.rename(temp_pdf_path, pdf_file)

        return os.path.exists(pdf_file)
    except (subprocess.SubprocessError, OSError):
        return False


def _convert_with_win32com(docx_file: str, pdf_file: str) -> bool:
    """Convert using Microsoft Word COM automation (Windows only)"""
    import platform

    if platform.system() != "Windows":
        return False

    try:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        doc = word.Documents.Open(os.path.abspath(docx_file))
        doc.SaveAs(os.path.abspath(pdf_file), FileFormat=17)  # 17 = PDF
        doc.Close()
        word.Quit()

        return os.path.exists(pdf_file)
    except ImportError:
        return False
    except Exception:
        # Clean up Word process if something went wrong
        try:
            doc.Close(False)
            word.Quit()
        except:
            pass
        return False


##############################
# Utility Helpers
##############################
def _url_image(
    img_para: DOCX_Paragraph,
    img_url: str,
    img_width: Inches,
    img_height: Inches,
    img_border_color: str = "000000",
    img_border_width: int = 1,
) -> DOCX_Run:
    """Add an image from a URL to a paragraph with optional border

    Args:
        img_para (DOCX_Paragraph): The paragraph to add the image to
        img_url (str): The URL of the image to add
        img_width (Inches): The width of the image
        img_height (Inches): The height of the image
        img_border_color (str): Hex color code for the image border (default: "000000")
        img_border_width (int): Width of the image border in points (default: 1)

    Returns:
        DOCX_Run: The run containing the image
    """
    run = img_para.add_run()

    try:
        from io import BytesIO

        import requests

        response = requests.get(img_url)
        img_stream = BytesIO(response.content)
        pic = run.add_picture(img_stream, width=img_width, height=img_height)

        # Get the inline element containing the picture
        inline = pic._inline

        # Alternative approach for handling shape properties that doesn't use namespaces parameter
        picPr = None
        for child in inline.findall(".//{*}pic"):
            picPr = child
            break

        if picPr is not None:
            # Find spPr element
            spPr = None
            for child in picPr.findall(".//{*}spPr"):
                spPr = child
                break

            if spPr is not None:
                # Create line properties element
                ln = OxmlElement("a:ln")
                ln.set("w", str(img_border_width * 12700))  # Width in EMUs

                # Create solid fill
                solidFill = OxmlElement("a:solidFill")
                srgbClr = OxmlElement("a:srgbClr")
                srgbClr.set("val", img_border_color.lstrip("#"))
                solidFill.append(srgbClr)
                ln.append(solidFill)

                # Add line properties to shape properties
                spPr.append(ln)
    except Exception as e:
        print(f"Failed to add image: {str(e)}")

    return run


def _paragraph_alignment(
    paragraph: DOCX_Paragraph, alignment_str: str = "center"
) -> None:
    """Set paragraph alignment based on string

    Args:
        paragraph: The paragraph to set alignment for
        alignment_str: Alignment as string (left, right, center)
    """
    alignment = alignment_str.lower()

    if alignment == "right":
        paragraph.alignment = DOCX_PARAGRAPH_ALIGN.RIGHT
    elif alignment == "left":
        paragraph.alignment = DOCX_PARAGRAPH_ALIGN.LEFT
    else:
        paragraph.alignment = DOCX_PARAGRAPH_ALIGN.CENTER


def _left_indent_paragraph(
    paragraph: DOCX_Paragraph, inches: float = 0.25
) -> DOCX_Paragraph:
    """Set the left indentation of a paragraph in inches

    Args:
        paragraph: The paragraph object to modify
        inches (float): Amount of indentation in inches

    Returns:
        paragraph: The modified paragraph object
    """
    paragraph.paragraph_format.left_indent = Inches(inches)
    return paragraph


def _process_list_item_formatting(
    paragraph: DOCX_Paragraph, li_element: BS4_Element, ensure_ending: bool = False
) -> None:
    """Process a list item and add its content to a paragraph with proper formatting

    Args:
        paragraph: The paragraph to add content to
        li_element: BeautifulSoup element for the list item
        ensure_ending: Whether to ensure the text ends with proper sentence ending
    """
    # Process the item content to preserve formatting
    for child in li_element.children:
        # Handle bold text (strong tags)
        if getattr(child, "name", None) == "strong":
            run = paragraph.add_run(child.text)
            run.bold = True
        # Handle italic text (em tags)
        elif getattr(child, "name", None) == "em":
            run = paragraph.add_run(child.text)
            run.italic = True
        # Handle links
        elif getattr(child, "name", None) == "a" and child.get("href"):
            _add_hyperlink(paragraph, child.text, child.get("href"))
        # Regular text
        elif child.string:
            text = child.string
            if ensure_ending:
                text = _ensure_sentence_ending(text)
            paragraph.add_run(text)


def _format_skills_list(
    document: DOCX_Document,
    skills: list[str],
    separator: str,
    apply_bold: bool = False,
    use_formatted_paragraph: bool = False,
) -> DOCX_Paragraph:
    """Add a skills list paragraph with optional bold formatting for individual skills

    Args:
        document: The Word document object
        skills: List of skill strings
        separator: Separator string between skills
        apply_bold: Whether to bold each individual skill
        use_formatted_paragraph: Whether to use _add_formatted_paragraph for non-bold case

    Returns:
        DOCX_Paragraph: The created paragraph
    """
    if apply_bold:
        # Create paragraph with bold skills and plain separators
        skills_para = document.add_paragraph()

        for i, skill in enumerate(skills):
            # Add separator before skills (except the first one)
            if i > 0:
                skills_para.add_run(separator)

            # Add the skill with bold formatting
            skill_run = skills_para.add_run(skill)
            skill_run.bold = True

        return skills_para
    else:
        # Use either _add_formatted_paragraph or create a simple paragraph
        if use_formatted_paragraph:
            return _add_formatted_paragraph(document, separator.join(skills))
        else:
            skills_para = document.add_paragraph()
            skills_para.add_run(separator.join(skills))
            return skills_para


def _add_formatted_paragraph(
    document: DOCX_Document,
    text: str,
    bold: bool = False,
    italic: bool = False,
    alignment: DOCX_PARAGRAPH_ALIGN = None,
    indentation: float = None,
    font_size: int | None = None,
) -> DOCX_Paragraph:
    """Add a paragraph with consistent formatting

    Args:
        document: The Word document object
        text (str): Text content for the paragraph
        bold (bool, optional): Whether text should be bold. Defaults to False.
        italic (bool, optional): Whether text should be italic. Defaults to False.
        alignment (DOCX_PARAGRAPH_ALIGN, optional): Paragraph alignment. Defaults to None.
        indentation (float, optional): Left indentation in inches. Defaults to None.
        font_size (int, optional): Font size in points. Defaults to None.

    Returns:
        paragraph: The created paragraph object
    """
    para = document.add_paragraph()

    # Check if text contains URLs, emails, or markdown links
    is_link = _detect_link(text)[0]

    if bold or italic or not is_link:
        # If bold/italic formatting is needed or no links detected, use simple formatting
        run = para.add_run(text)
        _apply_font_properties(
            run.font,
            {
                "bold": bold,
                "italic": italic,
                "font_size": font_size,
            },
        )
    else:
        # Process for hyperlinks
        _process_text_for_hyperlinks(para, text)

    # Apply paragraph-level formatting
    if alignment:
        para.alignment = alignment
    if indentation:
        _left_indent_paragraph(para, indentation)

    return para


def _has_hr_before_element(element: BS4_Element) -> bool:
    """Check if there's a horizontal rule (hr) element before any element

    Args:
        element: BeautifulSoup element to check for preceding HR

    Returns:
        bool: True if there's an HR element immediately before this element, False otherwise
    """
    prev = element.previous_sibling
    while prev and (not prev.name or prev.name == "br"):
        prev = prev.previous_sibling

    return prev and prev.name == "hr"


def _detect_link(text: str) -> tuple[bool, str, str, str]:
    """Detect if text contains any kind of link (markdown link, URL, or email)

    Args:
        text (str): Text to check

    Returns:
        tuple: (is_link, display_text, url, matched_text)
            - is_link (bool): Whether any link was found
            - display_text (str): Text to display for the link
            - url (str): The URL for the hyperlink
            - matched_text (str): The full text that matched (for extraction)
    """
    link_types = [
        {
            "pattern": MD_LINK_PATTERN,
            "formatter": lambda m: (m.group(1), m.group(2), m.group(0)),
        },
        {
            "pattern": URL_PATTERN,
            "formatter": lambda m: (m.group(0), _format_url(m.group(0)), m.group(0)),
        },
        {
            "pattern": EMAIL_PATTERN,
            "formatter": lambda m: (m.group(0), f"mailto:{m.group(0)}", m.group(0)),
        },
    ]

    for link_type in link_types:
        match = link_type["pattern"].search(text)
        if match:
            display_text, url, matched_text = link_type["formatter"](match)
            return True, display_text, url, matched_text

    return False, text, "", ""


def _format_url(url: str) -> str:
    """Format URL to ensure it has proper scheme

    Args:
        url (str): URL to format

    Returns:
        str: Formatted URL with scheme
    """
    if url.startswith("www."):
        return "http://" + url
    return url


def _process_text_for_hyperlinks(
    paragraph: DOCX_Paragraph, text: str, ensure_sentence_ending: bool = False
) -> None:
    """Process text to detect and add hyperlinks for Markdown links, URLs and email addresses

    Args:
        paragraph: The Word paragraph object to add content to
        text (str): Text to process for Markdown links, URLs and email addresses
        ensure_sentence_ending (bool): Whether to add period at the end if missing

    Returns:
        None: The paragraph is modified in place
    """
    # Check if text is None or empty
    if not text or not text.strip():
        return

    remaining_text = text
    fragments = []  # Store all text fragments and links

    # First pass: Identify all links and text segments
    while remaining_text:
        is_link, link_text, url, matched_text = _detect_link(remaining_text)

        if is_link:
            # Find the start position of the link
            link_position = remaining_text.find(matched_text)

            # Add text before the link as a text fragment
            if link_position > 0:
                fragments.append(("text", remaining_text[:link_position]))

            # Add the link as a link fragment
            fragments.append(("link", link_text, url))

            # Handle space after link
            after_link_pos = link_position + len(matched_text)
            if (
                after_link_pos < len(remaining_text)
                and remaining_text[after_link_pos] == " "
            ):
                fragments.append(("text", " "))
                after_link_pos += 1

            # Continue with remaining text
            remaining_text = remaining_text[after_link_pos:]
        else:
            # No more links, add all remaining text
            if remaining_text:
                fragments.append(("text", remaining_text))
            remaining_text = ""

    # Only add period to the very last text fragment if requested
    if ensure_sentence_ending and fragments and fragments[-1][0] == "text":
        last_idx = len(fragments) - 1
        text_content = fragments[last_idx][1]
        fragments[last_idx] = ("text", _ensure_sentence_ending(text_content))

    # Second pass: Add all fragments to paragraph
    for fragment in fragments:
        if fragment[0] == "text":
            paragraph.add_run(fragment[1])
        else:  # It's a link
            _, link_text, url = fragment
            _add_hyperlink(paragraph, link_text, url)


def _add_hyperlink(
    paragraph: DOCX_Paragraph, text: str, url: str
) -> docx.oxml.shared.OxmlElement:
    """Add a hyperlink to a paragraph using direct formatting instead of style reference

    Args:
        paragraph: The paragraph to add the hyperlink to
        text (str): The text to display for the hyperlink
        url (str): The URL to link to

    Returns:
        OxmlElement: The created hyperlink element
    """
    # Get access to the document
    part = paragraph.part
    # Create the relationship
    r_id = part.relate_to(url, DOCX_REL.HYPERLINK, is_external=True)

    # Create the hyperlink element
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)

    # Create a run inside the hyperlink
    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Create a run inside the hyperlink
    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Apply the Hyperlink style by referencing it
    style_id = docx.oxml.shared.OxmlElement("w:rStyle")
    style_id.set(docx.oxml.shared.qn("w:val"), "Hyperlink")
    rPr.append(style_id)

    # Add the run properties to the run
    new_run.append(rPr)
    # Set the text
    new_run.text = text
    # Add the run to the hyperlink
    hyperlink.append(new_run)

    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

    return hyperlink


def _add_horizontal_line_simple(document: DOCX_Document) -> None:
    """Add a simple horizontal line to the document using underscores

    Args:
        document: The Word document object

    Returns:
        None
    """
    # Skip adding horizontal line if two-column layout is enabled
    # if ConfigHelper.get_document_defaults().get("two_column_enabled", False):
    #     return

    # Get values from config with fallbacks
    line_char = ConfigHelper.get_style_constant("horizontal_line_char", "_")
    line_length = ConfigHelper.get_style_constant("horizontal_line_length", 50)
    line_color = ConfigHelper.get_style_constant("horizontal_line_color", "000000")

    p = document.add_paragraph()
    _paragraph_alignment(p, "center")
    p.add_run(line_char * line_length).bold = True

    # Set color for the horizontal line
    if p.runs:
        for run in p.runs:
            _apply_font_properties(
                run.font,
                {
                    "bold": True,
                    "color": line_color,
                },
            )


def _add_space_paragraph(
    document: DOCX_Document,
    font_size: int | None = None,
    space_before: int | None = None,
) -> None:
    """Add a paragraph with extra space after it

    Args:
        document: The Word document object
        font_size (int, optional): Controls spacing after paragraph and run font size.

    Returns:
        None
    """
    if font_size is None:
        font_size = ConfigHelper.get_style_constant("font_size_pts", 11)

    p = document.add_paragraph()
    p.add_run()
    # run.text = " "  # Add a space character to create a blank line
    # run.font.size = Pt(font_size)
    _add_space_before_or_after(p, space_before)


def _add_space_before_or_after(
    paragraph: DOCX_Paragraph,
    space_before: int | None = None,
    space_after: int | None = None,
) -> None:
    """Add a paragraph with extra space after it

    Args:
        document: The Word document object
        font_size (int, optional): Controls spacing after paragraph and run font size.

    Returns:
        None
    """
    if space_before:
        paragraph.paragraph_format.space_before = Pt(space_before)
    if space_after:
        paragraph.paragraph_format.space_after = Pt(space_after)


def _ensure_sentence_ending(text: str) -> str:
    """Ensures text ends with a period, question mark, or exclamation point

    Args:
        text (str): The text to check

    Returns:
        str: Text with proper sentence ending
    """
    if not text:
        return text

    # Strip trailing whitespace
    text = text.rstrip()

    # If already ends with sentence-ending punctuation, return as-is
    if text and text[-1] in [".", "!", "?", ":", ";"]:
        return text

    # Add a period
    return text + "."


def _find_organization_element(
    blockquote: BS4_Element,
) -> tuple[BS4_Element | None, bool]:
    """Find and return the organization element from a blockquote

    Args:
        blockquote: BeautifulSoup blockquote element

    Returns:
        tuple: (organization_element, was_processed)
    """
    for item in blockquote.contents:
        if isinstance(item, str) and not item.strip():
            continue

        if hasattr(item, "name") and item.name in ["h4", "h5", "h6"]:
            return item, False
        elif hasattr(item, "name") and item.name == "p":
            strong_tag = item.find("strong")
            if strong_tag:
                return item, False

    return None, False


def _create_hyperlink_style(
    document: DOCX_Document, hyperlink_props: Dict[str, str]
) -> bool:
    """Create a Hyperlink character style if it doesn't exist

    Args:
        document: The Word document object
        hyperlink_props: Dictionary containing hyperlink style properties

    Returns:
        bool: True if style was created successfully, False otherwise
    """
    if "Hyperlink" in document.styles:
        return True  # Already exists

    try:
        hyperlink_style = document.styles.add_style(
            "Hyperlink", DOCX_STYLE_TYPE.CHARACTER
        )
        validated_props = _validate_style_properties(hyperlink_props)
        _apply_font_properties(hyperlink_style.font, validated_props)
        return True
    except Exception as e:
        print(f"Warning: Could not create Hyperlink style: {str(e)}")
        return False


def _validate_style_properties(
    properties: Dict[str, str | int | float],
) -> Dict[str, str | int | float]:
    """Validate and clean style properties with type checking

    Args:
        properties: Raw properties dictionary

    Returns:
        dict: Validated and cleaned properties
    """
    valid_props = {}

    # Font properties with type validation
    font_prop_types = {
        "font_name": str,
        "font_size": (int, float),
        "bold": bool,
        "italic": bool,
        "underline": bool,
        "color": str,
    }

    for prop, expected_type in font_prop_types.items():
        if prop in properties:
            value = properties[prop]
            if isinstance(value, expected_type):
                valid_props[prop] = value
            else:
                print(
                    f"Warning: Invalid type for {prop}, expected {expected_type.__name__}, got {type(value).__name__}"
                )

    # Paragraph format properties with type validation
    para_prop_types = {
        "line_spacing": (int, float),
        "space_after": (int, float),
        "space_before": (int, float),
        "indent_left": (int, float),
        "indent_right": (int, float),
        "alignment": int,  # DOCX alignment constants
    }

    for prop, expected_type in para_prop_types.items():
        if prop in properties:
            value = properties[prop]
            if isinstance(value, expected_type):
                valid_props[prop] = value
            else:
                print(
                    f"Warning: Invalid type for {prop}, expected {expected_type.__name__}, got {type(value).__name__}"
                )

    return valid_props


def _apply_document_styles(
    document: DOCX_Document,
    styles: Dict[str, dict] = None,
    style_constants: Dict[str, str | dict] = None,
) -> None:
    """Apply style settings to document using values from config

    Args:
        document: The Word document object
        styles: Dictionary of style settings from config
        style_constants: Dictionary of style constants from config
    """
    if styles is None or style_constants is None:
        print("Warning: No styles or style constants provided, using defaults")
        return

    # Handle Hyperlink style creation if needed
    if "Hyperlink" in styles:
        _create_hyperlink_style(document, styles["Hyperlink"])

    # Apply properties to existing styles
    for style_name, properties in styles.items():
        if style_name == "Hyperlink" and style_name not in document.styles:
            continue  # Already handled above or creation failed

        try:
            style = document.styles[style_name]

            # Validate properties before applying them
            validated_props = _validate_style_properties(properties)

            # Apply font properties using helper function
            _apply_font_properties(style.font, validated_props)

            # Apply paragraph format properties if the style supports them
            if hasattr(style, "paragraph_format"):
                _apply_paragraph_format_properties(
                    style.paragraph_format, validated_props
                )

        except (KeyError, AttributeError, ValueError) as e:
            print(
                f"Warning: Could not apply all properties to '{style_name}': {str(e)}"
            )


__all__ = [
    "create_ats_resume",
    "convert_to_pdf",
    "ConfigLoader",
    "OutputFilePath",
    "DEFAULT_CONFIG_FILE",
    "DEFAULT_OUTPUT_DIR",
    "DEFAULT_OUTPUT_FORMAT",
    "DOCX_EXTENSION",
    "PDF_EXTENSION",
]

##############################
# Main Entry
##############################
if __name__ == "__main__":
    # Create detailed program description and examples
    program_description = """
    Convert a markdown resume to an ATS-friendly Word document.

    This script takes a markdown resume file and converts it to a properly formatted
    Word document that's optimized for Applicant Tracking Systems (ATS). It preserves
    the structure of your resume while ensuring proper formatting for employment history,
    skills, education, and other sections.
    """

    epilog_text = """
    Examples:
      python resume_md_to_docx.py
          - Runs in interactive mode (recommended for new users)

      python resume_md_to_docx.py -i resume.md
          - Converts resume.md to "output/resume.docx"

      python resume_md_to_docx.py -i resume.md -o resume.docx
          - Converts resume.md to resume.docx

      python resume_md_to_docx.py -i resume.md --pdf
          - Converts resume.md to "output/resume.docx" and "output/resume.pdf"

      python resume_md_to_docx.py -i resume.md -o resume.docx --pdf
          - Converts resume.md to resume.docx and resume.pdf
    """

    # Parse command line arguments with enhanced help
    parser = argparse.ArgumentParser(
        description=program_description,
        epilog=epilog_text,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    parser.add_argument("-i", "--input", dest="input_file", help="Input markdown file")
    parser.add_argument(
        "-c",
        "--config",
        dest="config_file",
        help="Path to YAML configuration file",
        default=DEFAULT_CONFIG_FILE,
    )
    parser.add_argument(
        "-o",
        "--output",
        dest="output_file",
        help='Output Word document (default: "<input_file>.docx")',
    )
    parser.add_argument(
        "-p",
        "--paragraph-headings",
        dest="paragraph_headings",
        nargs="+",
        choices=["h3", "h4", "h5", "h6"],
        help="Specify which heading levels to render as paragraphs instead of headings",
    )
    parser.add_argument(
        "-P",
        "--pdf",
        dest="create_pdf",
        action="store_true",
        help="Also create a PDF version of the resume",
    )
    parser.add_argument(
        "-I",
        "--interactive",
        dest="interactive",
        action="store_true",
        help="Run in interactive mode, prompting for inputs",
    )

    parser.add_argument(
        "-t",
        "--two-column",
        dest="two_column",
        action="store_true",
        help="Use two-column layout with tables",
    )

    args = parser.parse_args()

    # Check if we should run in interactive mode
    if (
        not (args.input_file or args.output_file or args.paragraph_headings)
        or args.interactive
    ):
        # Now we get the config_loader object directly from _run_interactive_mode
        input_file, output_file, paragraph_style_headings, create_pdf, config_loader = (
            _run_interactive_mode()
        )
    else:
        config_loader = ConfigLoader(args.config_file)

        if args.two_column:
            config_loader._config["document_defaults"]["two_column_enabled"] = True

        # Use command-line arguments
        input_file = Path(args.input_file)

        output_file = OutputFilePath(input_file, args.output_file).output_path()

        # Convert list to dictionary if provided
        paragraph_style_headings = {}
        if args.paragraph_headings:
            paragraph_style_headings = {h: True for h in args.paragraph_headings}

        # Show help if required arguments are missing
        if not input_file:
            parser.print_help()
            sys.exit(1)

        create_pdf = args.create_pdf

    # Use the arguments to create the resume, passing config_loader
    result = create_ats_resume(
        input_file,
        output_file,
        paragraph_style_headings=paragraph_style_headings,
        config_loader=config_loader,  # Pass config_loader object instead of string
    )

    # Convert to PDF if requested
    if create_pdf:
        pdf_file = convert_to_pdf(result)
        if pdf_file:
            print(f"✅ Created PDF: {pdf_file}")

    print(f"🎉 ATS-friendly resume created: {result} 🎉")
