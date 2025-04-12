import argparse
import os
import re
import sys
from enum import Enum
from pathlib import Path

import docx.oxml.shared
import markdown
from bs4 import BeautifulSoup
from bs4.element import PageElement as BS4_Element
from docx import Document
from docx.enum.style import WD_STYLE_TYPE as DOCX_STYLE_TYPE
from docx.enum.text import WD_BREAK as DOCX_PAGE_BREAK
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as DOCX_PARAGRAPH_ALIGN
from docx.opc.constants import RELATIONSHIP_TYPE as DOCX_REL
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph as DOCX_Paragraph

##############################
# Define default config file since it's referenced a few times
##############################
DEFAULT_CONFIG_FILE = "resume_config.yaml"

##############################
# Define some defaults at module level for better performance
##############################
MD_LINK_PATTERN = re.compile(r"\[(.*?)\]\((.*?)\)")
URL_PATTERN = re.compile(r"https?://[^\s]+|www\.[^\s]+")
EMAIL_PATTERN = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")


##############################
# Configs
##############################
class ResumeSection(Enum):
    """Maps markdown heading titles to their corresponding document headings

    Properties:
        markdown_heading (str): The text of the h2 heading in the markdown file
        docx_heading (str): The text to use as a heading in the Word document
        add_space_before_h3 (bool): Whether to add a blank line before each h3 heading
                                   (except the first after h2)
    """

    ABOUT = ("About", "PROFESSIONAL SUMMARY", False)
    SKILLS = ("Top Skills", "TOP SKILLS", False)
    EXPERIENCE = ("Experience", "PROFESSIONAL EXPERIENCE", True)
    EDUCATION = ("Education", "EDUCATION", True)
    CERTIFICATIONS = ("Licenses & certifications", "LICENSES & CERTIFICATIONS", True)
    CONTACT = ("Contact", "CONTACT INFORMATION", False)

    def __init__(
        self,
        markdown_heading: str,
        docx_heading: str,
        add_space_before_h3: bool = False,
    ):
        """Initialize ResumeSection enum

        Args:
            markdown_heading (str): The text of the h2 heading in the markdown file
            docx_heading (str): The text to use as a heading in the Word document
            add_space_before_h3 (bool): Whether to add a blank line before each h3 heading
                                       (except the first after h2). Defaults to False.
        """
        self.markdown_heading = markdown_heading
        self.docx_heading = docx_heading
        self.markdown_heading_lower = markdown_heading.lower()
        self.add_space_before_h3 = add_space_before_h3

    def matches(self, text):
        """Check if the given text matches this section's markdown_heading (case insensitive)

        Args:
            text (str): Text to compare against markdown_heading

        Returns:
            bool: True if text matches markdown_heading (case insensitive), False otherwise
        """
        return text.lower() == self.markdown_heading_lower


class JobSubsection(Enum):
    """Maps markdown subsection headings to their corresponding document headings

    Properties:
        markdown_heading_level (str): The HTML tag name (e.g., 'h5', 'h6') in markdown
        markdown_text_lower (str): The lowercase text content of the heading to match
        docx_heading (str): The text to use as a heading in the Word document
        separator (str): Optional separator to append after the heading (default: "")
        bold (bool): Whether the heading should be bold (default: True)
        italic (bool): Whether the heading should be italic (default: False)
    """

    KEY_SKILLS = ("h5", "key skills", "Key Skills", ":", True, False)
    SUMMARY = ("h5", "summary", "Summary", ":", True, False)
    INTERNAL = ("h5", "internal", "Internal", ":", True, False)
    PROJECT_CLIENT = ("h5", "project/client", "Project/Client", ": ", True, True)
    RESPONSIBILITIES = (
        "h6",
        "responsibilities overview",
        "Responsibilities",
        ":",
        True,
        False,
    )
    ADDITIONAL_DETAILS = (
        "h6",
        "additional details",
        "Additional Details",
        ":",
        True,
        False,
    )
    HIGHLIGHTS = ("h3", "highlights", "Highlights", "", True, False)

    def __init__(
        self,
        markdown_heading_level: str,
        markdown_text_lower: str,
        docx_heading: str,
        separator: str = "",
        bold: bool = True,
        italic: bool = False,
    ):
        """Initialize the JobSubsection with its properties"""
        self.markdown_heading_level = markdown_heading_level
        self.markdown_text_lower = markdown_text_lower
        self.docx_heading = docx_heading
        self.separator = separator
        self.bold = bold
        self.italic = italic

    @property
    def full_heading(self):
        """Return the full heading with separator

        Returns:
            str: The complete heading text with separator
        """
        return f"{self.docx_heading}{self.separator}"

    @classmethod
    def find_by_tag_and_text(cls, tag_name: str, text: str):
        """Find a JobSubsection by tag name and text content (case insensitive)

        Args:
            tag_name (str): HTML tag name to match (e.g., 'h5', 'h6')
            text (str): Text content to match (case insensitive)

        Returns:
            JobSubsection or None: The matching subsection or None if not found
        """
        text_lower = text.lower().strip()

        # First try exact match
        for subsection in cls:
            if (
                subsection.markdown_heading_level == tag_name
                and text_lower == subsection.markdown_text_lower
            ):
                return subsection

        # If no exact match, try partial match for elements
        # Create a dictionary of keywords to check against each subsection type
        if tag_name in [
            "h5",
            "h6",
        ]:  # Allow partial matching for both h5 and h6 elements
            # Group subsections by heading level
            matching_subsections = [
                s for s in cls if s.markdown_heading_level == tag_name
            ]

            # Try partial matching against each subsection's keywords
            for subsection in matching_subsections:
                # Extract the first word of the markdown_text_lower as the key keyword
                keyword = subsection.markdown_text_lower.split()[0]
                if keyword in text_lower:
                    return subsection

        return None


class PdfConverterPaths(Enum):
    """Enum to store paths to PDF converter executables on different platforms"""

    # Windows LibreOffice paths
    LIBREOFFICE_WINDOWS = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]

    # macOS LibreOffice path
    LIBREOFFICE_MACOS = "/Applications/LibreOffice.app/Contents/MacOS/soffice"

    # Linux/Unix LibreOffice command name (assumed to be in PATH)
    LIBREOFFICE_LINUX = "libreoffice"


##############################
# Helper Classes
##############################
class ConfigLoader:
    """Class for loading and accessing configuration from YAML file"""

    def __init__(
        self, config_file: Path = DEFAULT_CONFIG_FILE, print_success_msg: bool = True
    ):
        """Initialize by loading configuration from YAML file

        Args:
            config_file (Path): Path to configuration file.
                             Defaults to 'resume_config.yaml' in same directory.
            print_success_msg (bool): Whether to print success message after loading.
                                   Defaults to False.
        """
        import yaml
        from yaml.parser import ParserError

        # Default empty configuration structure
        self._config = {
            "document_defaults": {
                "output_filename": "My ATS Resume.docx",
                "pdf_file_extension": ".pdf",
                "margin_top_bottom": 0.7,
                "margin_left_right": 0.8,
                "page_width": 8.5,
                "page_height": 11.0,
            },
            "style_constants": {
                "font_size_pts": 11,
                "indent_inches": 0.25,
                "bullet_indent_inches": 0.5,
                "horizontal_line_char": "_",
                "horizontal_line_length": 50,
                "top_skills_separator": " | ",
                "bullet_separator": " • ",
                "paragraph_after_h4_line_spacing": 0.20,
            },
            "document_styles": {},
            "markdown_lists": {
                "ul": {
                    "bullet_character": "•",
                    "paragraph_delimiter": "\n",
                },
            },
            "markdown_headings": {
                "h1": {"level": 0, "paragraph_heading_size": 24},
                "h2": {"level": 1, "paragraph_heading_size": 16},
                "h3": {"level": 2, "paragraph_heading_size": 14},
                "h4": {"level": 3, "paragraph_heading_size": 12},
                "h5": {"level": 4, "paragraph_heading_size": 11},
                "h6": {"level": 5, "paragraph_heading_size": 10},
            },
        }

        # Try to load the YAML config file
        if os.path.exists(config_file):
            try:
                with open(config_file, "r") as f:
                    yaml_config = yaml.safe_load(f)

                # Use the YAML values directly if they exist
                if yaml_config and isinstance(yaml_config, dict):
                    # Replace document_defaults if provided
                    if "document_defaults" in yaml_config:
                        self._config["document_defaults"] = yaml_config[
                            "document_defaults"
                        ]

                    # Replace style_constants if provided
                    if "style_constants" in yaml_config:
                        self._config["style_constants"] = yaml_config["style_constants"]

                    if "markdown_lists" in yaml_config:
                        self._config["markdown_lists"] = yaml_config["markdown_lists"]

                    # Replace markdown_headings if provided
                    if "markdown_headings" in yaml_config:
                        self._config["markdown_headings"] = yaml_config[
                            "markdown_headings"
                        ]

                    # Process document styles - need to handle RGBColor and Pt objects
                    if "document_styles" in yaml_config:
                        document_styles = {}
                        for style_name, properties in yaml_config[
                            "document_styles"
                        ].items():
                            style_props = {}

                            # Convert font_size to Pt objects
                            if "font_size" in properties and isinstance(
                                properties["font_size"], (int, float)
                            ):
                                style_props["font_size"] = Pt(properties["font_size"])

                            # Convert color to RGBColor objects
                            if "color" in properties and isinstance(
                                properties["color"], str
                            ):
                                style_props["color"] = RGBColor.from_string(
                                    properties["color"]
                                )

                            # Convert indent values to Inches
                            for indent_key in ["indent_left", "indent_right"]:
                                if indent_key in properties and isinstance(
                                    properties[indent_key], (int, float)
                                ):
                                    style_props[indent_key] = Inches(
                                        properties[indent_key]
                                    )

                            # Copy over all other properties directly
                            for key, value in properties.items():
                                if key not in [
                                    "font_size",
                                    "color",
                                    "indent_left",
                                    "indent_right",
                                ]:
                                    style_props[key] = value

                            document_styles[style_name] = style_props

                        self._config["document_styles"] = document_styles

                    if print_success_msg:
                        print(f"✅ Config loaded from {config_file}")
            except (ParserError, Exception) as e:
                print(
                    f"❌ Error loading config file: {str(e)}, using default configuration"
                )

    @property
    def config(self) -> dict:
        """Get the entire configuration dictionary

        Returns:
            dict: Complete configuration dictionary
        """
        return self._config

    @property
    def document_defaults(self) -> dict:
        """Get document default settings

        Returns:
            dict: Document defaults configuration
        """
        return self._config.get("document_defaults", {})

    @property
    def style_constants(self) -> dict:
        """Get style constants

        Returns:
            dict: Style constants configuration
        """
        return self._config.get("style_constants", {})

    @property
    def document_styles(self) -> dict:
        """Get document styles

        Returns:
            dict: Document styles configuration
        """
        return self._config.get("document_styles", {})

    @property
    def markdown_headings(self) -> dict:
        """Get markdown headings configuration

        Returns:
            dict: Markdown headings configuration
        """
        return self._config.get("markdown_headings", {})

    def get_style_constant(self, key: str, default=None):
        """Get a specific style constant value

        Args:
            key (str): Style constant key to retrieve
            default: Default value if key is not found

        Returns:
            Value for the requested style constant or default if not found
        """
        return self.style_constants.get(key, default)


class ConfigHelper:
    """Static helper class for accessing configuration values globally"""

    # Class-level storage for configuration
    _config = None
    _initialized = False

    @classmethod
    def init(cls, config: dict) -> None:
        """Initialize configuration store

        Args:
            config (dict): Complete configuration dictionary
        """
        cls._config = config
        cls._initialized = True

    @classmethod
    def get_document_defaults(cls) -> dict:
        """Get document default settings

        Returns:
            dict: Document defaults configuration
        """
        cls._check_initialized()
        return cls._config.get("document_defaults", {})

    @classmethod
    def get_style_constants(cls) -> dict:
        """Get style constants

        Returns:
            dict: Style constants configuration
        """
        cls._check_initialized()
        return cls._config.get("style_constants", {})

    @classmethod
    def get_document_styles(cls) -> dict:
        """Get document styles

        Returns:
            dict: Document styles configuration
        """
        cls._check_initialized()
        return cls._config.get("document_styles", {})

    @classmethod
    def get_style_constant(cls, key: str, default=None):
        """Get a specific style constant value

        Args:
            key (str): Style constant key to retrieve
            default: Default value if key is not found

        Returns:
            Value for the requested style constant or default if not found
        """
        cls._check_initialized()
        return cls._config.get("style_constants", {}).get(key, default)

    @classmethod
    def get_markdown_list_option(cls, list_type: str, option_name: str, default=None):
        """Get a markdown list option from config

        Args:
            list_type: The list type ("ul" or "ol")
            option_name: The option name to retrieve
            default: Default value if not found

        Returns:
            The option value or default
        """
        cls._check_initialized()
        list_config = cls._config.get("markdown_lists", {})
        type_config = list_config.get(list_type, {})
        return type_config.get(option_name, default)

    @classmethod
    def _check_initialized(cls) -> None:
        """Check if the configuration has been initialized

        Raises:
            RuntimeError: If the configuration has not been initialized
        """
        if not cls._initialized:
            raise RuntimeError(
                "ConfigHelper has not been initialized. Call ConfigHelper.init() first."
            )


class HeadingsHelper:
    """Helper class for managing markdown headings based on configuration"""

    # Class-level storage for the heading configuration
    _heading_map = {}
    _paragraph_style_headings = {}
    _initialized = False

    @classmethod
    def init(
        cls, config: dict, paragraph_style_headings: dict[str, bool] = None
    ) -> None:
        """Initialize the heading map from configuration

        Args:
            config (dict): Configuration dictionary with markdown_headings
        """
        # Simply assign the markdown_headings from config
        cls._heading_map = config["markdown_headings"]
        cls._paragraph_style_headings = paragraph_style_headings or {}
        cls._initialized = True

    @classmethod
    def should_use_paragraph_style(cls, tag_name: str) -> bool:
        """Check if a heading tag should use paragraph style instead of heading style

        Args:
            tag_name (str): HTML tag name (e.g., 'h1', 'h2', etc.)

        Returns:
            bool: True if tag should use paragraph style, False otherwise
        """
        cls._check_initialized()
        return cls._paragraph_style_headings.get(tag_name, False)

    @classmethod
    def any_heading_uses_paragraph_style(cls) -> bool:
        """Check if any heading level uses paragraph style

        Returns:
            bool: True if any heading level has paragraph styling enabled, False otherwise
        """
        cls._check_initialized()
        return bool(cls._paragraph_style_headings)

    @classmethod
    def get_level_for_tag(cls, tag_name: str) -> int | None:
        """Get the Word document heading level for a given tag name

        Args:
            tag_name (str): HTML tag name (e.g., 'h1', 'h2', etc.)

        Returns:
            int or None: The corresponding Word document heading level or None if not found
        """
        cls._check_initialized()
        if tag_name.lower() in cls._heading_map:
            return cls._heading_map[tag_name.lower()]["level"]
        return None

    @classmethod
    def get_font_size_for_level(cls, heading_level: int, default_size: int = 11) -> int:
        """Get the font size for a given heading level

        Args:
            heading_level (int): The heading level (0-5)
            default_size (int): The default font size to return if not found

        Returns:
            int: The font size in points
        """
        cls._check_initialized()
        # Find the heading tag that corresponds to this level
        for tag, props in cls._heading_map.items():
            if props["level"] == heading_level:
                return props["paragraph_heading_size"]
        return default_size

    @classmethod
    def get_font_size_for_tag(cls, tag_name: str, default_size: int = 11) -> int:
        """Get the font size for a given tag name directly

        Args:
            tag_name (str): HTML tag name (e.g., 'h1', 'h2', etc.)
            default_size (int): The default font size to return if not found

        Returns:
            int: The font size in points for the tag
        """
        cls._check_initialized()
        tag_name = tag_name.lower()
        if tag_name in cls._heading_map:
            return cls._heading_map[tag_name]["paragraph_heading_size"]
        return default_size

    @classmethod
    def _check_initialized(cls) -> None:
        """Check if the configuration has been initialized

        Raises:
            RuntimeError: If the configuration has not been initialized
        """
        if not cls._initialized:
            raise RuntimeError(
                "ConfigHelper has not been initialized. Call ConfigHelper.init() first."
            )


##############################
# Main Processors
##############################
def create_ats_resume(
    md_file: Path,
    output_file: Path,
    config_loader: ConfigLoader,
    paragraph_style_headings: dict[str, bool] = None,
) -> Path:
    """Convert markdown resume to ATS-friendly Word document

    Args:
        md_file (Path): Path to the markdown resume file
        output_file (str): Path where the output Word document will be saved
        config_loader (ConfigLoader, optional): ConfigLoader instance with configuration.
                                               If None, creates with default config file.
        paragraph_style_headings (dict, optional): Dictionary mapping heading tags
                                                 to boolean values.

    Returns:
        Path: Path to the created document
    """
    # Default to using heading styles for all if not specified
    if paragraph_style_headings is None:
        paragraph_style_headings = {}

    # Initialize ConfigHelper and HeadingsHelper with config
    ConfigHelper.init(config_loader.config)
    HeadingsHelper.init(config_loader.config, paragraph_style_headings)

    # Access config sections directly through properties
    doc_defaults = config_loader.document_defaults
    style_constants = config_loader.style_constants
    document_styles = config_loader.document_styles

    # Read markdown file
    with open(md_file, "r") as file:
        md_content = file.read()

    # Convert markdown to HTML for easier parsing
    html = markdown.markdown(md_content)
    soup = BeautifulSoup(html, "html.parser")

    # Create document with standard margins
    document = Document()

    # Apply styles from configuration
    _apply_document_styles(document, document_styles, style_constants)

    for section in document.sections:
        section.page_width = Inches(doc_defaults["page_width"])
        section.page_height = Inches(doc_defaults["page_height"])
        section.top_margin = Inches(doc_defaults["margin_top_bottom"])
        section.bottom_margin = Inches(doc_defaults["margin_top_bottom"])
        section.left_margin = Inches(doc_defaults["margin_left_right"])
        section.right_margin = Inches(doc_defaults["margin_left_right"])

    # Define section processors with their specific processing functions
    section_processors = [
        (ResumeSection.ABOUT, process_header_section),
        (
            ResumeSection.ABOUT,
            lambda doc, soup: process_about_section(document=doc, soup=soup),
        ),
        (
            ResumeSection.SKILLS,
            lambda doc, soup: process_skills_section(doc, soup),
        ),
        (
            ResumeSection.EXPERIENCE,
            lambda doc, soup: process_experience_section(doc, soup),
        ),
        (
            ResumeSection.EDUCATION,
            lambda doc, soup: process_education_section(doc, soup),
        ),
        (
            ResumeSection.CERTIFICATIONS,
            lambda doc, soup: process_certifications_section(doc, soup),
        ),
        (
            ResumeSection.CONTACT,
            lambda doc, soup: process_contact_section(doc, soup),
        ),
    ]

    # Process each section
    for section_type, processor in section_processors:
        processor(document, soup)

    # Save the document
    document.save(output_file)

    return Path(output_file)


def convert_to_pdf(docx_file: Path) -> Path | None:
    """Convert a DOCX file to PDF using available converters

    Args:
        docx_file (Path): Path to the input DOCX file

    Returns:
        Path: Path to the created PDF file, or None if conversion failed
    """
    pdf_file = docx_file.with_suffix(".pdf")

    # Try multiple conversion methods
    methods = [
        _convert_with_docx2pdf,
        _convert_with_libreoffice,
        _convert_with_win32com,
    ]

    for method in methods:
        try:
            if method(docx_file, pdf_file):
                return Path(pdf_file)
        except Exception as e:
            print(f"Could not convert using {method.__name__}: {str(e)}")

    print("❌ PDF conversion failed. Please install one of the following:")
    print("   - docx2pdf (pip install docx2pdf)")
    print("   - LibreOffice (https://www.libreoffice.org/)")
    print("   - Microsoft Word (Windows only)")
    return None


##############################
# Section Processors
##############################
def process_header_section(
    document: Document,
    soup: BeautifulSoup,
) -> None:
    """Process the header (name and tagline) section

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content

    Returns:
        None
    """
    # Extract header (name)
    name = soup.find("h1").text

    # Add name as document title
    title = document.add_heading(name, HeadingsHelper.get_level_for_tag("h1"))
    title.alignment = DOCX_PARAGRAPH_ALIGN.CENTER

    # Add professional tagline if it exists - first paragraph after h1
    first_p = soup.find("h1").find_next_sibling()
    if first_p and first_p.name == "p":
        # Check if ANY paragraph headings are specified, which indicates preference for simpler styling
        use_paragraph_style = (
            HeadingsHelper.any_heading_uses_paragraph_style()
        )  # True if any heading levels use paragraph style

        # Check if the paragraph contains emphasis (italics)
        em_tag = first_p.find("em")
        if em_tag:
            # Always use paragraph style with manual formatting when paragraph_style_headings is active
            if use_paragraph_style:
                tagline_para = document.add_paragraph()
                tagline_para.alignment = DOCX_PARAGRAPH_ALIGN.CENTER
                tagline_run = tagline_para.add_run(em_tag.text)
                tagline_run.italic = True

                # Set the font size to match heading style
                tagline_run.font.size = Pt(HeadingsHelper.get_font_size_for_tag("h4"))
            else:
                # Use Word's built-in Subtitle style
                tagline_para = document.add_paragraph(em_tag.text, style="Subtitle")
                tagline_para.alignment = DOCX_PARAGRAPH_ALIGN.CENTER
                # Keep it italic despite the style
                for run in tagline_para.runs:
                    run.italic = True

            # Add the rest of the first paragraph as a separate paragraph if it exists
            rest_of_p = first_p.text.replace(em_tag.text, "").strip()
            if rest_of_p:
                rest_para = document.add_paragraph()
                rest_para.alignment = DOCX_PARAGRAPH_ALIGN.CENTER
                rest_para.add_run(rest_of_p)
        else:
            # If no emphasis tag, just add the whole paragraph
            if use_paragraph_style:
                tagline_para = document.add_paragraph()
                tagline_para.alignment = DOCX_PARAGRAPH_ALIGN.CENTER
                tagline_para.add_run(first_p.text)
            else:
                # Use Word's built-in Subtitle style
                tagline_para = document.add_paragraph(first_p.text, style="Subtitle")
                tagline_para.alignment = DOCX_PARAGRAPH_ALIGN.CENTER

        # Check for additional paragraphs before the first h2 that might contain specialty areas
        current_p = first_p.find_next_sibling()
        # Simply process all paragraphs until we hit a non-paragraph element (like h2)
        while current_p and current_p.name == "p":
            specialty_para = document.add_paragraph()
            specialty_para.alignment = DOCX_PARAGRAPH_ALIGN.CENTER
            specialty_para.add_run(current_p.text)
            current_p = current_p.find_next_sibling()

    # Add horizontal line
    _add_horizontal_line_simple(document)


def process_about_section(
    document: Document,
    soup: BeautifulSoup,
) -> None:
    """Process the About section

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content

    Returns:
        None
    """
    section_h2 = _prepare_section(document, soup, ResumeSection.ABOUT)

    if not section_h2:
        return

    current_element = section_h2.find_next_sibling()

    # Initialize processed_elements set if we're in the About section
    processed_elements = set()

    # Add all paragraphs until next h2
    while current_element and current_element.name != "h2":
        # Skip if already processed
        if current_element in processed_elements:
            current_element = current_element.find_next_sibling()
            continue

        # Check if this is a paragraph with strong element containing highlights
        highlights_subsection = None
        if (
            current_element.name == JobSubsection.HIGHLIGHTS.markdown_heading_level
            and current_element.text.strip().lower()
            == JobSubsection.HIGHLIGHTS.markdown_text_lower
        ):
            highlights_subsection = JobSubsection.HIGHLIGHTS

        # Handle regular paragraph
        if not highlights_subsection and current_element.name == "p":
            para = document.add_paragraph()

            # Process all elements of the paragraph to preserve formatting
            for child in current_element.children:
                # Check if this is a strong/bold element
                if getattr(child, "name", None) == "strong":
                    run = para.add_run(child.text)
                    run.bold = True
                # Check if this is an em/italic element
                elif getattr(child, "name", None) == "em":
                    run = para.add_run(child.text)
                    run.italic = True
                # Check if this is a link/anchor element
                elif getattr(child, "name", None) == "a" and child.get("href"):
                    _add_hyperlink(para, child.text, child.get("href"))
                # Otherwise, just add the text as-is
                else:
                    if child.string:
                        _process_text_for_hyperlinks(para, child.string)

            processed_elements.add(current_element)

        # Handle highlights subsection found in a paragraph or heading
        elif highlights_subsection:
            heading_level = (
                HeadingsHelper.get_level_for_tag(current_element.name)
                if current_element.name.startswith("h")
                else None
            )
            use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
                current_element.name
            )
            _add_heading_or_paragraph(
                document,
                highlights_subsection.full_heading,
                heading_level,
                use_paragraph_style=use_paragraph_style,
                bold=highlights_subsection.bold,
                italic=highlights_subsection.italic,
            )
            processed_elements.add(current_element)

        # Handle bullet list
        elif current_element.name == "ul":
            _add_bullet_list(document, current_element)
            processed_elements.add(current_element)

        current_element = current_element.find_next_sibling()


def process_skills_section(
    document: Document,
    soup: BeautifulSoup,
) -> None:
    """Process the Skills section

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content

    Returns:
        None
    """
    section_h2 = _prepare_section(document, soup, ResumeSection.SKILLS)

    if not section_h2:
        return

    current_element = section_h2.find_next_sibling()

    if current_element and current_element.name == "p":
        input_separator = ConfigHelper.get_style_constant(
            "top_skills_separator_markdown", " • "
        )
        skills = [
            s.strip() for s in current_element.text.split(input_separator.strip())
        ]
        skills = [s for s in skills if s]
        output_separator = ConfigHelper.get_style_constant(
            "top_skills_separator", " | "
        )
        _add_formatted_paragraph(document, output_separator.join(skills))


def process_experience_section(
    document: Document,
    soup: BeautifulSoup,
) -> None:
    """Process the Experience section

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content

    Returns:
        None
    """
    section_h2 = _prepare_section(document, soup, ResumeSection.EXPERIENCE)

    if not section_h2:
        return

    # Find all job entries (h3 headings under Experience)
    current_element = section_h2.find_next_sibling()
    # Use a set of element IDs instead of element objects
    processed_element_ids = set()
    processed_elements = set()
    add_space_before_h3 = ResumeSection.EXPERIENCE.add_space_before_h3

    while current_element and current_element.name != "h2":
        # Get unique ID for this element
        element_id = id(current_element)

        # Skip if already processed (except h3 elements)
        if element_id in processed_element_ids and current_element.name != "h3":
            current_element = current_element.find_next_sibling()
            continue

        # Process based on element type
        if current_element.name == "h3":
            # Process job entry and mark as processed
            processed_elements = _process_job_entry(
                document, current_element, set(), add_space_before_h3
            )
            processed_element_ids.add(element_id)
        elif current_element.name == "h4" and element_id not in processed_element_ids:
            # Process position and mark as processed
            processed_elements = _process_position(document, current_element, set())
            processed_element_ids.add(element_id)

        elif (
            current_element.name in ["h5", "h6"]
            and element_id not in processed_element_ids
        ):
            # Find matching subsection type
            subsection = JobSubsection.find_by_tag_and_text(
                current_element.name, current_element.text
            )

            if subsection:
                heading_level = HeadingsHelper.get_level_for_tag(current_element.name)

                # PROJECT/CLIENT requires special handling with its own function
                if (
                    subsection == JobSubsection.PROJECT_CLIENT
                    and current_element not in processed_elements
                ):
                    project_processed = _process_project_section(
                        document,
                        current_element,
                        processed_elements,
                    )

                    # Add all elements processed by the project handler to our tracking
                    processed_elements.update(project_processed)

                    # Also update element IDs
                    for element in project_processed:
                        processed_element_ids.add(id(element))

                # Generic subsection processing for SUMMARY, INTERNAL, etc.
                elif (
                    subsection in [JobSubsection.SUMMARY, JobSubsection.INTERNAL]
                    and current_element not in processed_elements
                ):
                    project_processed = _process_subsection(
                        document,
                        current_element,
                        subsection,
                        heading_level,
                        processed_elements,
                    )

                    # Update tracking
                    processed_elements.update(project_processed)
                    for element in project_processed:
                        processed_element_ids.add(id(element))

                # KEY_SKILLS subsection
                elif subsection == JobSubsection.KEY_SKILLS:
                    use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
                        current_element.name
                    )
                    skills_heading = _add_heading_or_paragraph(
                        document,
                        subsection.full_heading,
                        heading_level,
                        use_paragraph_style=use_paragraph_style,
                        bold=subsection.bold,
                        italic=subsection.italic,
                    )
                    skills_para = document.add_paragraph()

                    # Get skills from next element
                    next_element = current_element.find_next_sibling()
                    if next_element:
                        if next_element.name == "p":
                            input_separator = ConfigHelper.get_style_constant(
                                "key_skills_separator_markdown", " • "
                            )
                            skills = [
                                s.strip()
                                for s in next_element.text.split(
                                    input_separator.strip()
                                )
                            ]
                            skills = [s for s in skills if s]
                            output_separator = ConfigHelper.get_style_constant(
                                "key_skills_separator", " • "
                            )
                            skills_para.add_run(output_separator.join(skills))
                            processed_elements.add(next_element)

                    # Determine if we need to add a blank line after Key Skills
                    # We'll add a blank line if:
                    # 1. There are no more headings (end of section)
                    # 2. The next heading is an h3 (new job)
                    # 3. The next heading is an h4 (new role within same company)
                    # We won't add a blank line if there's an h5 or h6 after Key Skills
                    looking_ahead = current_element
                    next_heading = None

                    # Look for the next heading element
                    while looking_ahead and not next_heading:
                        looking_ahead = looking_ahead.find_next_sibling()
                        if looking_ahead and looking_ahead.name in [
                            "h3",
                            "h4",
                            "h5",
                            "h6",
                        ]:
                            next_heading = looking_ahead

                    # Add space if this is the last role or before a new role (most likely h4)
                    if not next_heading or next_heading.name in ["h4"]:
                        _add_space_paragraph(document, 8)

                # RESPONSIBILITIES subsection (standalone)
                elif (
                    subsection == JobSubsection.RESPONSIBILITIES
                    and current_element not in processed_elements
                ):
                    use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
                        current_element.name
                    )
                    _add_heading_or_paragraph(
                        document,
                        subsection.full_heading,
                        heading_level,
                        use_paragraph_style=use_paragraph_style,
                        bold=subsection.bold,
                        italic=subsection.italic,
                    )

                    # Get content
                    next_element = current_element.find_next_sibling()
                    if next_element:
                        if next_element.name == "p":
                            resp_para = document.add_paragraph()
                            _process_text_for_hyperlinks(resp_para, next_element.text)
                            processed_elements.add(next_element)
                        elif next_element.name == "ul":
                            # Process bullet list
                            _add_bullet_list(document, next_element)
                            processed_elements.add(next_element)

                # ADDITIONAL_DETAILS subsection (standalone)
                elif (
                    subsection == JobSubsection.ADDITIONAL_DETAILS
                    and current_element not in processed_elements
                ):
                    use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
                        current_element.name
                    )
                    _add_heading_or_paragraph(
                        document,
                        subsection.full_heading,
                        heading_level,
                        use_paragraph_style=use_paragraph_style,
                        bold=subsection.bold,
                        italic=subsection.italic,
                    )

                    # Get content (next element might be list items)
                    next_element = current_element.find_next_sibling()
                    if next_element and next_element.name == "ul":
                        _add_bullet_list(document, next_element)
                        processed_elements.add(next_element)

        # Standalone bullet points
        elif current_element.name == "ul" and current_element not in processed_elements:
            # Process bullet list
            _add_bullet_list(document, current_element)
            processed_element_ids.add(element_id)

        current_element = current_element.find_next_sibling()


def process_education_section(
    document: Document,
    soup: BeautifulSoup,
) -> None:
    """Process the Education section

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content

    Returns:
        None
    """
    section_h2 = _prepare_section(
        document,
        soup,
        ResumeSection.EDUCATION,
    )
    _process_simple_section(
        document,
        section_h2,
        add_space=ResumeSection.EDUCATION.add_space_before_h3,
    )


def process_certifications_section(
    document: Document,
    soup: BeautifulSoup,
) -> None:
    """Process the Certifications section

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content

    Returns:
        None
    """
    section_h2 = _prepare_section(
        document,
        soup,
        ResumeSection.CERTIFICATIONS,
    )
    _process_certifications(
        document,
        section_h2,
    )


def process_contact_section(
    document: Document,
    soup: BeautifulSoup,
) -> None:
    """Process the Contact section

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content

    Returns:
        None
    """
    section_h2 = _prepare_section(
        document,
        soup,
        ResumeSection.CONTACT,
    )
    _process_simple_section(
        document,
        section_h2,
        add_space=ResumeSection.CONTACT.add_space_before_h3,
    )


##############################
# Primary Helpers
##############################
def _apply_document_styles(
    document: Document, styles: dict = None, style_constants: dict = None
) -> None:
    """Apply style settings to document using values from config

    Args:
        document: The Word document object
        styles: Dictionary of style settings from config
        style_constants: Dictionary of style constants from config
    """
    if styles is None or style_constants is None:
        print("Warning: No styles or style constants provided, using defaults")
        return

    if "Hyperlink" in styles and "Hyperlink" not in document.styles:
        hyperlink_props = styles["Hyperlink"]
        try:
            hyperlink_style = document.styles.add_style(
                "Hyperlink", DOCX_STYLE_TYPE.CHARACTER
            )

            # Apply font properties from the config
            if "font_name" in hyperlink_props:
                hyperlink_style.font.name = hyperlink_props["font_name"]
            if "font_size" in hyperlink_props:
                hyperlink_style.font.size = hyperlink_props["font_size"]
            if "underline" in hyperlink_props:
                hyperlink_style.font.underline = hyperlink_props["underline"]
            if "color" in hyperlink_props and isinstance(hyperlink_props["color"], str):
                hyperlink_style.font.color.rgb = RGBColor.from_string(
                    hyperlink_props["color"]
                )
        except Exception as e:
            print(f"Warning: Could not create Hyperlink style: {str(e)}")

    for style_name, properties in styles.items():

        if style_name == "Hyperlink" and style_name not in document.styles:
            continue  # Already handled above

        try:
            # Get the style object from document
            style = document.styles[style_name]

            # Apply font properties
            if "font_name" in properties:
                style.font.name = properties["font_name"]
            if "font_size" in properties:
                style.font.size = properties["font_size"]
            if "bold" in properties:
                style.font.bold = properties["bold"]
            if "italic" in properties:
                style.font.italic = properties["italic"]
            if "underline" in properties:
                style.font.underline = properties["underline"]
            if "color" in properties:
                style.font.color.rgb = properties["color"]

            # Apply paragraph format properties
            if hasattr(style, "paragraph_format"):
                if "line_spacing" in properties:
                    style.paragraph_format.line_spacing = properties["line_spacing"]
                if "space_after" in properties:
                    style.paragraph_format.space_after = properties["space_after"]
                if "indent_left" in properties:
                    style.paragraph_format.left_indent = properties["indent_left"]
                if "indent_right" in properties:
                    style.paragraph_format.right_indent = properties["indent_right"]
                if "alignment" in properties:
                    style.paragraph_format.alignment = properties["alignment"]

        except (KeyError, AttributeError) as e:
            print(
                f"Warning: Could not apply all properties to '{style_name}': {str(e)}"
            )


def _prepare_section(
    document: Document,
    soup: BeautifulSoup,
    section_type: ResumeSection,
) -> BS4_Element | None:
    """Universal preliminary section preparation

    Args:
        document: The Word document object
        soup: BeautifulSoup object of the HTML content
        section_type: ResumeSection enum value

    Returns:
        BeautifulSoup element or None: The section heading element if found, None otherwise
    """
    section_h2 = soup.find("h2", string=lambda text: section_type.matches(text))

    if not section_h2:
        return None

    # Use the generic function instead of the specific one
    section_page_break = _has_hr_before_element(section_h2)

    # Add page break if requested
    if section_page_break:
        p = document.add_paragraph()
        run = p.add_run()
        run.add_break(DOCX_PAGE_BREAK.PAGE)

    # Add the section heading
    use_paragraph_style = HeadingsHelper.should_use_paragraph_style("h2")
    heading_level = HeadingsHelper.get_level_for_tag("h2")
    _add_heading_or_paragraph(
        document,
        section_type.docx_heading,
        heading_level,
        use_paragraph_style=use_paragraph_style,
    )

    return section_h2


def _process_simple_section(
    document: Document,
    section_h2: BS4_Element,
    add_space: bool = False,
) -> None:
    """Process sections with simple paragraph-based content like Education and Contact.
    These sections typically have paragraphs with some bold (strong) elements.

    Args:
        document: The Word document object
        section_h2: The BeautifulSoup h2 element for the section
        add_space (bool, optional): Whether to add a space paragraph after the section. Defaults to False.

    Returns:
        None
    """
    current_element = section_h2.find_next_sibling()

    while current_element and current_element.name != "h2":
        if current_element.name == "p":
            para = document.add_paragraph()
            for child in current_element.children:
                if getattr(child, "name", None) == "strong":
                    para.add_run(f"{child.text}: ").bold = True
                else:
                    if child.string and child.string.strip():
                        _process_text_for_hyperlinks(para, child.string.strip())
        elif current_element.name == "ul":
            # Handle bullet lists if they appear
            _add_bullet_list(document, current_element)

        current_element = current_element.find_next_sibling()

    # Add an extra space after the section if requested
    if add_space:
        _add_space_paragraph(document, ConfigHelper.get_style_constant("font_size_pts"))


def _process_project_section(
    document: Document,
    project_element: BS4_Element,
    processed_elements: set[BS4_Element],
) -> set[BS4_Element]:
    """Process a project/client section and its related elements

    Args:
        document: The Word document object
        project_element: BeautifulSoup element for the project heading
        processed_elements: Set of elements already processed

    Returns:
        set: Updated set of processed elements
    """
    bullet_indent_inches = ConfigHelper.get_style_constant("bullet_indent_inches")

    # Get the next element to see if it contains the project details
    next_element = project_element.find_next_sibling()
    project_info = ""

    # If next element is a paragraph, it might contain the project name/duration
    if next_element and next_element.name == "p":
        project_info = next_element.text.strip()
        processed_elements.add(next_element)
        next_element = next_element.find_next_sibling()

    # Get the proper subsection
    subsection = JobSubsection.find_by_tag_and_text(
        project_element.name, project_element.text
    )
    if not subsection:
        subsection = JobSubsection.PROJECT_CLIENT  # Fallback

    heading_level = HeadingsHelper.get_level_for_tag(project_element.name)
    use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
        project_element.name
    )

    # Prepare the project text
    project_text = subsection.full_heading
    if project_info:
        project_text += " " + project_info

    # Add the heading or formatted paragraph
    _add_heading_or_paragraph(
        document,
        project_text,
        heading_level,
        use_paragraph_style=use_paragraph_style,
        bold=subsection.bold,
        italic=subsection.italic,
    )

    # Process next elements under this project until another section
    while (
        next_element
        and next_element.name not in ["h2", "h3", "h4"]
        and not (next_element.name == "h5")
    ):
        if next_element.name == "p" and not next_element.find("h6"):
            # Process regular paragraph text
            para = document.add_paragraph()
            _process_text_for_hyperlinks(para, next_element.text.strip())
            _left_indent_paragraph(para, bullet_indent_inches / 2)
            processed_elements.add(next_element)
            next_element = next_element.find_next_sibling()
            continue

        # Find subsection for h6 elements
        h6_subsection = None
        if next_element.name == "h6":
            h6_subsection = JobSubsection.find_by_tag_and_text(
                next_element.name, next_element.text
            )

        # Responsibilities Overview
        if h6_subsection == JobSubsection.RESPONSIBILITIES:
            heading_level = HeadingsHelper.get_level_for_tag(next_element.name)
            use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
                next_element.name
            )

            # Add the heading or paragraph
            if use_paragraph_style:
                resp_header = document.add_paragraph()
                _left_indent_paragraph(resp_header)  # Keep indentation
                resp_run = resp_header.add_run(h6_subsection.full_heading)
                resp_run.bold = h6_subsection.bold
                resp_run.italic = h6_subsection.italic
            else:
                resp_heading = document.add_heading(
                    h6_subsection.full_heading, level=heading_level
                )
                _left_indent_paragraph(resp_heading)  # Keep indentation

            # Get the paragraph with responsibilities
            resp_element = next_element.find_next_sibling()
            if resp_element and resp_element.name == "p":
                resp_para = document.add_paragraph()
                _process_text_for_hyperlinks(resp_para, resp_element.text)
                _left_indent_paragraph(resp_para)
                processed_elements.add(resp_element)

            processed_elements.add(next_element)

        # Additional Details
        elif h6_subsection == JobSubsection.ADDITIONAL_DETAILS:
            heading_level = HeadingsHelper.get_level_for_tag(next_element.name)
            use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
                next_element.name
            )

            # Add the heading or paragraph
            if use_paragraph_style:
                details_header = document.add_paragraph()
                _left_indent_paragraph(details_header)  # Keep indentation
                details_run = details_header.add_run(h6_subsection.full_heading)
                details_run.bold = h6_subsection.bold
                details_run.italic = h6_subsection.italic
            else:
                details_heading = document.add_heading(
                    h6_subsection.full_heading, level=heading_level
                )
                _left_indent_paragraph(details_heading)  # Keep indentation

            processed_elements.add(next_element)

        # Bullet points
        elif next_element.name == "ul":
            _left_indent_paragraph(_add_bullet_list(document, next_element))
            processed_elements.add(next_element)

        next_element = next_element.find_next_sibling()

    return processed_elements


def _process_job_entry(
    document: Document,
    job_element: BS4_Element,
    processed_elements: set[BS4_Element],
    add_space_before_h3: bool = False,
) -> set[BS4_Element]:
    """Process a job entry (h3) and its related elements

    Args:
        document: The Word document object
        job_element: BeautifulSoup element for the job heading
        processed_elements: Set of elements already processed
        add_space_before_h3: Whether to add space before h3 headings (except first)

    Returns:
        Updated set of processed elements
    """
    job_title = job_element.text.strip()

    # Check for HR before h3 and add page break if found
    if _has_hr_before_element(job_element):
        p = document.add_paragraph()
        run = p.add_run()
        run.add_break(DOCX_PAGE_BREAK.PAGE)
    # Otherwise add normal spacing if needed
    elif add_space_before_h3:
        prev_heading = job_element.find_previous(["h2", "h3"])
        if prev_heading and prev_heading.name == "h3":
            document.add_paragraph()

    # Add the company name as h3
    use_paragraph_style = HeadingsHelper.should_use_paragraph_style("h3")
    heading_level = HeadingsHelper.get_level_for_tag("h3")
    _add_heading_or_paragraph(
        document, job_title, heading_level, use_paragraph_style=use_paragraph_style
    )

    # Mark the h3 as processed
    processed_elements.add(job_element)

    # Check for the paragraph immediately after h3
    next_element = job_element.find_next_sibling()

    # Process paragraph with duration text if it exists
    if next_element and next_element.name == "p":
        duration_para = document.add_paragraph()

        # Check for both bold and italic formatting
        has_strong = next_element.find("strong")
        has_em = next_element.find("em")

        duration_text = next_element.text.replace("*", "").replace("_", "").strip()
        duration_run = duration_para.add_run(duration_text)

        if has_strong:
            duration_run.bold = True
        if has_em:
            duration_run.italic = True

        processed_elements.add(next_element)

    return processed_elements


def _process_subsection(
    document: Document,
    current_element: BS4_Element,
    subsection: JobSubsection,
    heading_level: int,
    processed_elements: set[BS4_Element],
) -> set[BS4_Element]:
    """Generic function to process any subsection (Summary, Internal, Responsibilities, etc.)

    Args:
        document: The Word document object
        current_element: The subsection heading element
        subsection: The JobSubsection enum value
        heading_level: The heading level from HeadingsHelper
        processed_elements: Set of elements already processed

    Returns:
        Updated set of processed elements
    """
    if current_element not in processed_elements:

        processed_elements.add(current_element)

        # Add the subsection heading
        use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
            current_element.name
        )
        _add_heading_or_paragraph(
            document,
            subsection.full_heading,
            heading_level,
            use_paragraph_style=use_paragraph_style,
            bold=subsection.bold,
            italic=subsection.italic,
        )

        # Process elements under this subsection until we hit another heading
        next_element = current_element.find_next_sibling()
        stop_tags = ["h2", "h3", "h4", "h5", "h6"]

        while next_element and next_element.name not in stop_tags:
            if next_element.name == "p":
                para = document.add_paragraph()
                _process_text_for_hyperlinks(para, next_element.text.strip())
                processed_elements.add(next_element)
            elif next_element.name == "ul":
                _add_bullet_list(document, next_element)
                processed_elements.add(next_element)

            next_element = next_element.find_next_sibling()

    return processed_elements


def _process_position(
    document: Document,
    element: BS4_Element,
    processed_elements: set[BS4_Element],
) -> set[BS4_Element]:
    """Process a position entry (h4) and its related elements

    Args:
        document: The Word document object
        job: BeautifulSoup element
        processed_elements: Set of elements already processed

    Returns:
        Updated set of processed elements
    """
    position_title = element.text.strip()

    # Add the position heading
    heading_level = HeadingsHelper.get_level_for_tag(element.name)
    use_paragraph_style = HeadingsHelper.should_use_paragraph_style(element.name)
    _add_heading_or_paragraph(
        document,
        position_title,
        heading_level,
        use_paragraph_style=use_paragraph_style,
    )

    processed_elements.add(element)
    next_element = element.find_next_sibling()

    # Process date and location elements (either h6 or p format)
    if next_element:
        # Handle paragraph date/location format
        # Handle paragraph with combined date and location on single line
        if next_element.name == "p":
            if next_element.find("strong") or next_element.find("em"):

                # Check if next paragraph is location
                next_next_element = next_element.find_next_sibling()
                if (
                    next_next_element
                    and next_next_element.name == "p"
                    and (
                        next_next_element.find("strong") or next_next_element.find("em")
                    )
                ):
                    # Create a single paragraph with both date and location
                    date_text = (
                        next_element.text.replace("*", "").replace("_", "").strip()
                    )
                    location_text = (
                        next_next_element.text.replace("*", "").replace("_", "").strip()
                    )

                    reduced_spacing = float(
                        ConfigHelper.get_style_constant(
                            "paragraph_after_h4_line_spacing", 0.80
                        )
                    )
                    combo_para = document.add_paragraph()
                    combo_para.paragraph_format.line_spacing = reduced_spacing

                    # Get font size from config
                    date_loc_font_size = ConfigHelper.get_style_constant(
                        "paragraph_after_h4_font_size", None
                    )

                    # Add date with appropriate formatting
                    date_run = combo_para.add_run(date_text + " - ")
                    if next_element.find("strong"):
                        date_run.bold = True
                    if next_element.find("em"):
                        date_run.italic = True
                    if date_loc_font_size:
                        date_run.font.size = Pt(date_loc_font_size)

                    # Add location with appropriate formatting
                    location_run = combo_para.add_run(location_text)
                    if next_next_element.find("strong"):
                        location_run.bold = True
                    if next_next_element.find("em"):
                        location_run.italic = True
                    if date_loc_font_size:
                        location_run.font.size = Pt(date_loc_font_size)

                    processed_elements.add(next_element)
                    processed_elements.add(next_next_element)
                    next_element = next_next_element.find_next_sibling()

    return processed_elements


def _add_heading_or_paragraph(
    document: Document,
    text: str,
    heading_level: int,
    use_paragraph_style: bool = False,
    bold: bool = True,
    italic: bool = False,
    font_size: int = None,
) -> DOCX_Paragraph:
    """Add either a heading or a formatted paragraph based on preference

    Args:
        document: The Word document object
        text (str): Text content for the heading/paragraph
        heading_level (int): The heading level (0-5) to use if not using paragraph style
        use_paragraph_style (bool): Whether to use paragraph style instead of heading style
        bold (bool): Whether to make the paragraph text bold (if using paragraph style)
        italic (bool): Whether to make the paragraph text italic (if using paragraph style)
        font_size (int, optional): Font size in points for paragraph style. Defaults to None.

    Returns:
        The created heading or paragraph object
    """
    if use_paragraph_style:
        para = document.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic

        # Apply appropriate font size from HeadingsHelper if not explicitly provided
        if font_size is None:
            size_pt = HeadingsHelper.get_font_size_for_level(heading_level)
            run.font.size = Pt(size_pt)
        else:
            run.font.size = Pt(font_size)

        return para
    else:
        return document.add_heading(text, level=heading_level)


def _process_certifications(
    document: Document,
    section_h2: BeautifulSoup,
) -> None:
    """Process the certifications section with its specific structure

    Args:
        document: The Word document object
        section_h2: BeautifulSoup h2 element for the section

    Returns:
        None
    """
    # Get the add_space_before_h3 setting from the section type
    add_space_before_h3 = ResumeSection.CERTIFICATIONS.add_space_before_h3

    current_element = section_h2.find_next_sibling()
    first_h3_after_h2 = True  # Track the first h3 after h2

    while current_element and current_element.name != "h2":
        # Process certification name (h3)
        if current_element.name == "h3":
            cert_name = current_element.text.strip()
            use_paragraph_style = HeadingsHelper.should_use_paragraph_style("h3")
            heading_level = HeadingsHelper.get_level_for_tag("h3")

            # Add blank line before h3 except for the first one
            # Only add space if add_space_before_h3 is True
            if add_space_before_h3 and not first_h3_after_h2:
                document.add_paragraph()
            else:
                first_h3_after_h2 = False  # After first h3 is processed

            _add_heading_or_paragraph(
                document,
                cert_name,
                heading_level,
                use_paragraph_style=use_paragraph_style,
            )

            # Look for next elements - either blockquote or organization info directly
            next_element = current_element.find_next_sibling()

            # Handle blockquote (optional)
            if next_element and next_element.name == "blockquote":
                # Process the blockquote contents
                _process_certification_blockquote(document, next_element)

            # If no blockquote, look for organization info directly
            elif next_element:
                # Try to find organization info (could be bold text or heading)
                if next_element.name in ["h4", "h5", "h6", "p"] and next_element.find(
                    "strong"
                ):
                    # Extract organization text
                    org_text = next_element.find("strong").text.strip()
                    org_para = document.add_paragraph()
                    org_run = org_para.add_run(org_text)
                    org_run.bold = True

                    # Look for date information in the next element
                    date_element = next_element.find_next_sibling()
                    if date_element and date_element.find("em"):
                        em_tag = date_element.find("em")
                        date_text = em_tag.text.strip()
                        date_para = document.add_paragraph()

                        # Check if the date is hyperlinked
                        parent_a = em_tag.find_parent("a")
                        if parent_a and parent_a.get("href"):
                            _add_hyperlink(date_para, date_text, parent_a["href"])
                        else:
                            date_run = date_para.add_run(date_text)
                            date_run.italic = True

            # Add spacing after each certification
            # _add_space_paragraph(document)

        current_element = current_element.find_next_sibling()


def _process_certification_blockquote(
    document: Document,
    blockquote: BS4_Element,
) -> None:
    """Process the contents of a certification blockquote

    Args:
        document: The Word document object
        blockquote: BeautifulSoup blockquote element containing certification details

    Returns:
        None
    """
    # Process the organization name (first element)
    org_element = blockquote.find(["h4", "h5", "h6", "strong"])
    if org_element:
        if org_element.name.startswith("h"):
            # It's a heading
            heading_level = HeadingsHelper.get_level_for_tag(org_element.name)
            use_paragraph_style = HeadingsHelper.should_use_paragraph_style(
                org_element.name
            )
            _add_heading_or_paragraph(
                document,
                org_element.text.strip(),
                heading_level,
                use_paragraph_style=use_paragraph_style,
            )
        elif org_element.name == "strong":
            # It's bold text
            org_para = document.add_paragraph()
            org_run = org_para.add_run(org_element.text.strip())
            org_run.bold = True

    # Process all paragraphs and text content
    # Find all p tags and loose text nodes directly under blockquote
    for item in blockquote.contents:
        # Skip elements we've already processed or that are empty
        if (isinstance(item, str) and not item.strip()) or item == org_element:
            continue

        # Check if it's a heading (h5, h6, etc.)
        if (
            hasattr(item, "name")
            and item.name
            and item.name.startswith("h")
            and item != org_element
        ):
            heading_level = HeadingsHelper.get_level_for_tag(item.name)
            use_paragraph_style = HeadingsHelper.should_use_paragraph_style(item.name)
            _add_heading_or_paragraph(
                document,
                item.text.strip(),
                heading_level,
                use_paragraph_style=use_paragraph_style,
            )
            continue

        # Check if it's the date with italics (should be last non-empty element)
        if hasattr(item, "find") and item.find("em"):
            em_tag = item.find("em")
            date_text = em_tag.text.strip()
            date_para = document.add_paragraph()

            # Check if inside a hyperlink
            parent_a = em_tag.find_parent("a")
            if parent_a and parent_a.get("href"):
                _add_hyperlink(date_para, date_text, parent_a["href"])
            else:
                date_run = date_para.add_run(date_text)
                date_run.italic = True
            continue

        # Handle normal text content (like "Some details")
        if isinstance(item, str) and item.strip():
            # It's a direct text node
            para = document.add_paragraph()
            _process_text_for_hyperlinks(para, item.strip())
        elif hasattr(item, "name"):
            if item.name == "p":
                # It's a paragraph
                para = document.add_paragraph()
                _process_text_for_hyperlinks(para, item.text.strip())
            elif item.name == "ul":
                # It's a bullet list
                _add_bullet_list(document, item)


##############################
# Inractive Mode Helper
##############################
def _run_interactive_mode() -> tuple[str, str, dict[str, bool], bool, ConfigLoader]:
    """Run in interactive mode, prompting the user for inputs

    Returns:
        tuple: (input_file, output_file, paragraph_style_headings, create_pdf, config_loader)
    """
    print("\n🎯 Welcome to Resume Markdown to ATS Converter (Interactive Mode) 🎯\n")

    # Prompt for config file first
    default_config = DEFAULT_CONFIG_FILE
    config_prompt = (
        f"⚙️ Enter path to configuration file (default: '{default_config}'): "
    )
    config_file = input(config_prompt).strip()
    if not config_file:
        config_file = default_config
        print(f"✅ Using default configuration: {config_file}")

    # Create ConfigLoader with the specified config file
    config_loader = ConfigLoader(config_file)

    # Prompt for input file
    while True:
        input_file = input("📄 Enter the path to your Markdown resume file: ").strip()
        if not input_file:
            print("❌ Input file path cannot be empty. Please try again.")
            continue

        if not os.path.exists(input_file):
            print(f"❌ File '{input_file}' does not exist. Please enter a valid path.")
            continue

        break

    # Prompt for output file
    default_output = config_loader.document_defaults["output_filename"]
    output_prompt = f"📝 Enter the output docx filename (default: '{default_output}'): "
    output_file = input(output_prompt).strip()
    if not output_file:
        output_file = default_output
        print(f"✅ Using default output: {output_file}")

    # Prompt for paragraph style headings
    print(
        "\n🔠 Choose heading levels to render as paragraphs instead of Word headings:"
    )
    print("   (Enter numbers separated by space, e.g., '3 4 5 6' for h3, h4, h5, h6)")
    print("   1. h3 - Job titles")
    print("   2. h4 - Company names")
    print("   3. h5 - Subsections (Key Skills, Summary, etc.)")
    print("   4. h6 - Sub-subsections (Responsibilities, Additional Details)")
    print("   0. None (use Word heading styles for all)")

    heading_choices = input(
        "👉 Your choices (e.g., '3 4 5 6' or '0' for none): "
    ).strip()

    paragraph_style_headings = {}
    if heading_choices != "0":
        chosen_numbers = [int(n) for n in heading_choices.split() if n.isdigit()]
        heading_map = {1: "h3", 2: "h4", 3: "h5", 4: "h6"}

        for num in chosen_numbers:
            if 1 <= num <= 4:
                heading_tag = heading_map[num]
                paragraph_style_headings[heading_tag] = True
                print(f"✅ Selected {heading_tag} for paragraph styling")
    else:
        print("✅ Using Word headings for all levels (no paragraph styling)")

    print("\n⚙️ Processing your resume...\n")

    create_pdf = (
        input("📄 Also create a PDF version? (y/n, default: n): ").strip().lower()
        == "y"
    )
    if create_pdf:
        print("✅ Will generate PDF output")

    # Return the ConfigLoader object directly instead of just the config_file path
    return input_file, output_file, paragraph_style_headings, create_pdf, config_loader


##############################
# PDF Helpers
##############################
def _convert_with_docx2pdf(docx_file: str, pdf_file: str) -> bool:
    """Convert using docx2pdf library"""
    try:
        from docx2pdf import convert

        convert(docx_file, pdf_file)
        return os.path.exists(pdf_file)
    except ImportError:
        return False


def _convert_with_libreoffice(docx_file: str, pdf_file: str) -> bool:
    """Convert using LibreOffice command line"""
    import platform
    import subprocess

    # Find the LibreOffice executable based on platform
    if platform.system() == "Windows":
        for path in PdfConverterPaths.LIBREOFFICE_WINDOWS.value:
            if os.path.exists(path):
                lo_exec = path
                break
        else:
            return False
    elif platform.system() == "Darwin":  # macOS
        lo_exec = PdfConverterPaths.LIBREOFFICE_MACOS.value
        if not os.path.exists(lo_exec):
            return False
    else:  # Linux/Unix
        lo_exec = PdfConverterPaths.LIBREOFFICE_LINUX.value  # Assume it's in PATH

    try:
        subprocess.run(
            [
                lo_exec,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                os.path.dirname(pdf_file) or ".",
                docx_file,
            ],
            check=True,
            stdout=subprocess.DEVNULL,
        )

        # LibreOffice saves to the original filename with .pdf extension
        temp_pdf = os.path.splitext(os.path.basename(docx_file))[0] + ".pdf"
        temp_pdf_path = os.path.join(os.path.dirname(pdf_file) or ".", temp_pdf)

        # If the output path is different from LibreOffice's default, move the file
        if os.path.abspath(temp_pdf_path) != os.path.abspath(pdf_file):
            os.rename(temp_pdf_path, pdf_file)

        return os.path.exists(pdf_file)
    except (subprocess.SubprocessError, OSError):
        return False


def _convert_with_win32com(docx_file: str, pdf_file: str) -> bool:
    """Convert using Microsoft Word COM automation (Windows only)"""
    import platform

    if platform.system() != "Windows":
        return False

    try:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        doc = word.Documents.Open(os.path.abspath(docx_file))
        doc.SaveAs(os.path.abspath(pdf_file), FileFormat=17)  # 17 = PDF
        doc.Close()
        word.Quit()

        return os.path.exists(pdf_file)
    except ImportError:
        return False
    except Exception:
        # Clean up Word process if something went wrong
        try:
            doc.Close(False)
            word.Quit()
        except:
            pass
        return False


##############################
# Utilities
##############################
def _left_indent_paragraph(
    paragraph: DOCX_Paragraph, inches: float = 0.25
) -> DOCX_Paragraph:
    """Set the left indentation of a paragraph in inches

    Args:
        paragraph: The paragraph object to modify
        inches (float): Amount of indentation in inches

    Returns:
        paragraph: The modified paragraph object
    """
    # paragraph.paragraph_format._left_indent_paragraph = Inches(inches)
    paragraph.paragraph_format.left_indent = Inches(inches)
    return paragraph


def _add_bullet_list(
    document: Document, ul_element: BS4_Element, indentation: float = None
) -> DOCX_Paragraph:
    """Add bullet points from an unordered list element

    Args:
        document: The Word document object
        ul_element: BeautifulSoup element containing the unordered list
        indentation (float, optional): Left indentation in inches

    Returns:
        paragraph: The last bullet paragraph added
    """
    use_paragraph_style = ConfigHelper.get_style_constant("paragraph_lists", False)

    if use_paragraph_style:
        # Get the bullet character from config
        bullet_char = ConfigHelper.get_markdown_list_option(
            "ul", "bullet_character", "•"
        )

        para = document.add_paragraph()

        items = ul_element.find_all("li")
        for i, li in enumerate(items):
            if i > 0:
                para.add_run("\n")

            # Add the bullet character first
            para.add_run(f"{bullet_char} ")

            # Process formatting using the helper function
            # Always ensure sentence ending for each line in paragraph lists
            _process_list_item_formatting(para, li, ensure_ending=True)

        if indentation:
            _left_indent_paragraph(para, indentation)

        return para
    else:
        # Standard bullet list processing - unchanged
        bullet_para = None
        for li in ul_element.find_all("li"):
            bullet_para = document.add_paragraph(style="List Bullet")

            # Process formatting using the helper function
            _process_list_item_formatting(bullet_para, li)

            # Ensure sentence ending for each bullet
            last_run = bullet_para.runs[-1] if bullet_para.runs else None
            if last_run and not last_run.text.rstrip()[-1:] in [
                ".",
                "!",
                "?",
                ":",
                ";",
            ]:
                last_run.text = last_run.text.rstrip() + "."

            if indentation:
                _left_indent_paragraph(bullet_para, indentation)

        return bullet_para


def _process_list_item_formatting(
    paragraph: DOCX_Paragraph, li_element: BS4_Element, ensure_ending: bool = False
) -> None:
    """Process a list item and add its content to a paragraph with proper formatting

    Args:
        paragraph: The paragraph to add content to
        li_element: BeautifulSoup element for the list item
        ensure_ending: Whether to ensure the text ends with proper sentence ending
    """
    # Process the item content to preserve formatting
    for child in li_element.children:
        # Handle bold text (strong tags)
        if getattr(child, "name", None) == "strong":
            run = paragraph.add_run(child.text)
            run.bold = True
        # Handle italic text (em tags)
        elif getattr(child, "name", None) == "em":
            run = paragraph.add_run(child.text)
            run.italic = True
        # Handle links
        elif getattr(child, "name", None) == "a" and child.get("href"):
            _add_hyperlink(paragraph, child.text, child.get("href"))
        # Regular text
        elif child.string:
            text = child.string
            if ensure_ending:
                text = _ensure_sentence_ending(text)
            paragraph.add_run(text)


def _add_formatted_paragraph(
    document: Document,
    text: str,
    bold: bool = False,
    italic: bool = False,
    alignment: DOCX_PARAGRAPH_ALIGN = None,
    indentation: float = None,
    font_size: int = None,
) -> DOCX_Paragraph:
    """Add a paragraph with consistent formatting

    Args:
        document: The Word document object
        text (str): Text content for the paragraph
        bold (bool, optional): Whether text should be bold. Defaults to False.
        italic (bool, optional): Whether text should be italic. Defaults to False.
        alignment (DOCX_PARAGRAPH_ALIGN, optional): Paragraph alignment. Defaults to None.
        indentation (float, optional): Left indentation in inches. Defaults to None.
        font_size (int, optional): Font size in points. Defaults to None.

    Returns:
        paragraph: The created paragraph object
    """
    para = document.add_paragraph()

    # Check if text contains URLs, emails, or markdown links
    is_link = _detect_link(text)[0]

    if bold or italic or not is_link:
        # If bold/italic formatting is needed or no links detected, use simple formatting
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        if font_size:
            run.font.size = Pt(font_size)
    else:
        # Process for hyperlinks
        _process_text_for_hyperlinks(para, text)

    # Apply paragraph-level formatting
    if alignment:
        para.alignment = alignment
    if indentation:
        _left_indent_paragraph(para, indentation)

    return para


def _has_hr_before_element(element: BS4_Element) -> bool:
    """Check if there's a horizontal rule (hr) element before any element

    Args:
        element: BeautifulSoup element to check for preceding HR

    Returns:
        bool: True if there's an HR element immediately before this element, False otherwise
    """
    if not element:
        return False

    prev_element = element.previous_sibling
    # Skip whitespace text nodes
    while prev_element and isinstance(prev_element, str) and prev_element.strip() == "":
        prev_element = prev_element.previous_sibling

    # Check if the previous element is an HR
    return prev_element and prev_element.name == "hr"


def _detect_link(text: str) -> tuple[bool, str, str, str]:
    """Detect if text contains any kind of link (markdown link, URL, or email)

    Args:
        text (str): Text to check

    Returns:
        tuple: (is_link, display_text, url, matched_text)
            - is_link (bool): Whether any link was found
            - display_text (str): Text to display for the link
            - url (str): The URL for the hyperlink
            - matched_text (str): The full text that matched (for extraction)
    """
    link_types = [
        {
            "pattern": MD_LINK_PATTERN,
            "formatter": lambda m: (m.group(1), m.group(2), m.group(0)),
        },
        {
            "pattern": URL_PATTERN,
            "formatter": lambda m: (m.group(0), _format_url(m.group(0)), m.group(0)),
        },
        {
            "pattern": EMAIL_PATTERN,
            "formatter": lambda m: (m.group(0), f"mailto:{m.group(0)}", m.group(0)),
        },
    ]

    for link_type in link_types:
        match = link_type["pattern"].search(text)
        if match:
            display_text, url, matched_text = link_type["formatter"](match)
            return True, display_text, url, matched_text

    return False, text, "", ""


def _format_url(url: str) -> str:
    """Format URL to ensure it has proper scheme

    Args:
        url (str): URL to format

    Returns:
        str: Formatted URL with scheme
    """
    if url.startswith("www."):
        return "http://" + url
    return url


def _process_text_for_hyperlinks(
    paragraph: DOCX_Paragraph, text: str, ensure_sentence_ending: bool = False
) -> None:
    """Process text to detect and add hyperlinks for Markdown links, URLs and email addresses

    Args:
        paragraph: The Word paragraph object to add content to
        text (str): Text to process for Markdown links, URLs and email addresses
        ensure_sentence_ending (bool): Whether to add period at the end if missing

    Returns:
        None: The paragraph is modified in place
    """
    # Check if text is None or empty
    if not text or not text.strip():
        return

    remaining_text = text
    fragments = []  # Store all text fragments and links

    # First pass: Identify all links and text segments
    while remaining_text:
        is_link, link_text, url, matched_text = _detect_link(remaining_text)

        if is_link:
            # Find the start position of the link
            link_position = remaining_text.find(matched_text)

            # Add text before the link as a text fragment
            if link_position > 0:
                fragments.append(("text", remaining_text[:link_position]))

            # Add the link as a link fragment
            fragments.append(("link", link_text, url))

            # Handle space after link
            after_link_pos = link_position + len(matched_text)
            if (
                after_link_pos < len(remaining_text)
                and remaining_text[after_link_pos] == " "
            ):
                fragments.append(("text", " "))
                after_link_pos += 1

            # Continue with remaining text
            remaining_text = remaining_text[after_link_pos:]
        else:
            # No more links, add all remaining text
            if remaining_text:
                fragments.append(("text", remaining_text))
            remaining_text = ""

    # Only add period to the very last text fragment if requested
    if ensure_sentence_ending and fragments and fragments[-1][0] == "text":
        last_idx = len(fragments) - 1
        text_content = fragments[last_idx][1]
        fragments[last_idx] = ("text", _ensure_sentence_ending(text_content))

    # Second pass: Add all fragments to paragraph
    for fragment in fragments:
        if fragment[0] == "text":
            paragraph.add_run(fragment[1])
        else:  # It's a link
            _, link_text, url = fragment
            _add_hyperlink(paragraph, link_text, url)


def _add_hyperlink(
    paragraph: DOCX_Paragraph, text: str, url: str
) -> docx.oxml.shared.OxmlElement:
    """Add a hyperlink to a paragraph using direct formatting instead of style reference

    Args:
        paragraph: The paragraph to add the hyperlink to
        text (str): The text to display for the hyperlink
        url (str): The URL to link to

    Returns:
        OxmlElement: The created hyperlink element
    """
    # Get access to the document
    part = paragraph.part
    # Create the relationship
    r_id = part.relate_to(url, DOCX_REL.HYPERLINK, is_external=True)

    # Create the hyperlink element
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)

    # Create a run inside the hyperlink
    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Create a run inside the hyperlink
    new_run = docx.oxml.shared.OxmlElement("w:r")
    rPr = docx.oxml.shared.OxmlElement("w:rPr")

    # Apply the Hyperlink style by referencing it
    style_id = docx.oxml.shared.OxmlElement("w:rStyle")
    style_id.set(docx.oxml.shared.qn("w:val"), "Hyperlink")
    rPr.append(style_id)

    # Add the run properties to the run
    new_run.append(rPr)
    # Set the text
    new_run.text = text
    # Add the run to the hyperlink
    hyperlink.append(new_run)

    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

    return hyperlink


def _add_horizontal_line_simple(document: Document) -> None:
    """Add a simple horizontal line to the document using underscores

    Args:
        document: The Word document object

    Returns:
        None
    """
    # Get values from config with fallbacks
    line_char = ConfigHelper.get_style_constant("horizontal_line_char", "_")
    line_length = ConfigHelper.get_style_constant("horizontal_line_length", 50)

    p = document.add_paragraph()
    p.alignment = DOCX_PARAGRAPH_ALIGN.CENTER
    p.add_run(line_char * line_length).bold = True


def _add_space_paragraph(document: Document, font_size: int = None) -> None:
    """Add a paragraph with extra space after it

    Args:
        document: The Word document object
        font_size (int, optional): Controls spacing after paragraph and run font size.

    Returns:
        None
    """
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(font_size)


def _ensure_sentence_ending(text: str) -> str:
    """Ensures text ends with a period, question mark, or exclamation point

    Args:
        text (str): The text to check

    Returns:
        str: Text with proper sentence ending
    """
    if not text:
        return text

    # Strip trailing whitespace
    text = text.rstrip()

    # If already ends with sentence-ending punctuation, return as-is
    if text and text[-1] in [".", "!", "?", ":", ";"]:
        return text

    # Add a period
    return text + "."


##############################
# Main Entry
##############################
if __name__ == "__main__":
    # Create detailed program description and examples
    program_description = """
    Convert a markdown resume to an ATS-friendly Word document.

    This script takes a markdown resume file and converts it to a properly formatted
    Word document that's optimized for Applicant Tracking Systems (ATS). It preserves
    the structure of your resume while ensuring proper formatting for employment history,
    skills, education, and other sections.
    """

    epilog_text = """
    Examples:
      python resume_md_to_docx.py
          - Runs in interactive mode (recommended for new users)

      python resume_md_to_docx.py -i resume.md
          - Converts resume.md to "My ATS Resume.docx"

      python resume_md_to_docx.py -i resume.md -o resume.docx
          - Converts resume.md to resume.docx

      python resume_md_to_docx.py -i resume.md --pdf
          - Converts resume.md to "My ATS Resume.docx" and "My ATS Resume.pdf"

      python resume_md_to_docx.py -i resume.md -o resume.docx --pdf
          - Converts resume.md to resume.docx and resume.pdf
    """

    # Parse command line arguments with enhanced help
    parser = argparse.ArgumentParser(
        description=program_description,
        epilog=epilog_text,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    parser.add_argument("-i", "--input", dest="input_file", help="Input markdown file")
    parser.add_argument(
        "-c",
        "--config",
        dest="config_file",
        help="Path to YAML configuration file",
        default=DEFAULT_CONFIG_FILE,
    )
    parser.add_argument(
        "-o",
        "--output",
        dest="output_file",
        help='Output Word document (default: "My ATS Resume.docx")',
    )
    parser.add_argument(
        "-p",
        "--paragraph-headings",
        dest="paragraph_headings",
        nargs="+",
        choices=["h3", "h4", "h5", "h6"],
        help="Specify which heading levels to render as paragraphs instead of headings",
    )
    parser.add_argument(
        "-P",
        "--pdf",
        dest="create_pdf",
        action="store_true",
        help="Also create a PDF version of the resume",
    )
    parser.add_argument(
        "-I",
        "--interactive",
        dest="interactive",
        action="store_true",
        help="Run in interactive mode, prompting for inputs",
    )

    args = parser.parse_args()

    # Check if we should run in interactive mode
    if (
        not (args.input_file or args.output_file or args.paragraph_headings)
        or args.interactive
    ):
        # Now we get the config_loader object directly from _run_interactive_mode
        input_file, output_file, paragraph_style_headings, create_pdf, config_loader = (
            _run_interactive_mode()
        )
    else:
        config_loader = ConfigLoader(args.config_file)

        # Use command-line arguments
        input_file = args.input_file
        output_file = (
            args.output_file or config_loader.document_defaults["output_filename"]
        )

        # Convert list to dictionary if provided
        paragraph_style_headings = {}
        if args.paragraph_headings:
            paragraph_style_headings = {h: True for h in args.paragraph_headings}

        # Show help if required arguments are missing
        if not input_file:
            parser.print_help()
            sys.exit(1)

        create_pdf = args.create_pdf

    # Use the arguments to create the resume, passing config_loader
    result = create_ats_resume(
        input_file,
        output_file,
        paragraph_style_headings=paragraph_style_headings,
        config_loader=config_loader,  # Pass config_loader object instead of string
    )

    # Convert to PDF if requested
    if create_pdf:
        pdf_file = convert_to_pdf(result)
        if pdf_file:
            print(f"✅ Created PDF: {pdf_file}")

    print(f"🎉 ATS-friendly resume created: {result} 🎉")
